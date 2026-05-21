"""
AI-POC 详细设计文档生成脚本
使用 python-docx 生成标准格式的 Word 文档

详细设计文档聚焦：类/方法级签名、算法流程、状态机、事件协议、错误码、测试点
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = Path(__file__).parent / "AI-POC详细设计文档.docx"


# ================= 样式辅助 =================

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:color'), '808080')
        tcBorders.append(tag)


def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def h(doc, text, level=1):
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, name='黑体', size=sizes.get(level, 12), bold=True, color=(0x1F, 0x3A, 0x5F))
    return p


def para(doc, text, indent=True, bold=False, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_font(run, name='宋体', size=size, bold=bold)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    for r in p.runs:
        r.text = ''
    run = p.add_run(text)
    set_font(run, name='宋体', size=11)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, name='Consolas', size=9, color=(0x2F, 0x4F, 0x4F))
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F5F5F5')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    hdr_cells = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(head)
        set_font(run, name='黑体', size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F3A5F')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        set_cell_border(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_data in rows:
        row_cells = t.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ''
            pc = row_cells[i].paragraphs[0]
            pc.paragraph_format.line_spacing = 1.3
            run = pc.add_run(str(val))
            set_font(run, name='宋体', size=9)
            set_cell_border(row_cells[i])
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t


# ================= 构建文档 =================

def build():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # ---------- 封面 ----------
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI-POC 漏洞验证平台')
    set_font(run, name='黑体', size=32, bold=True, color=(0x1F, 0x3A, 0x5F))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('详细设计文档')
    set_font(run, name='黑体', size=26, bold=True, color=(0x1F, 0x3A, 0x5F))

    for _ in range(8):
        doc.add_paragraph()

    meta = [
        ('文档名称', 'AI-POC 漏洞验证平台详细设计文档'),
        ('文档版本', 'V1.0'),
        ('编写日期', '2026-04-18'),
        ('文档状态', '初稿'),
        ('适用范围', '详细设计阶段 / 编码实现指导'),
        ('前置文档', 'AI-POC概要设计文档 V1.0'),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, (k, v) in enumerate(meta):
        row = t.rows[i]
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
        for j, text in enumerate((k, v)):
            cell = row.cells[j]
            cell.text = ''
            pc = cell.paragraphs[0]
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pc.add_run(text)
            set_font(run, name='宋体', size=12, bold=(j == 0))
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    # ---------- 目录 ----------
    h(doc, '目  录', level=1)
    toc = [
        '1 引言', '2 设计约束与全局依赖', '3 业务服务层详细设计',
        '4 数据模型与数据库详细设计', '5 接口详细设计',
        '6 前端详细设计', '7 POC 脚本规范',
        '8 关键流程时序设计', '9 错误处理与日志',
        '10 性能与并发设计', '11 测试设计', '12 附录',
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.6
        run = p.add_run(item)
        set_font(run, name='宋体', size=12)

    doc.add_page_break()

    # ========== 1 引言 ==========
    h(doc, '1 引言', level=1)

    h(doc, '1.1 编写目的', level=2)
    para(doc, '本文档是 AI-POC 漏洞验证平台的详细设计文档，在《AI-POC 概要设计文档》确定系统架构、模块划分、主要流程的基础上，进一步给出类与方法的接口签名、内部算法、数据结构、状态机、事件协议与错误码等实现级细节。本文档的目标读者是编码与测试阶段的开发人员，阅读后应当能够不依赖口头沟通即可完成对应模块的实现或替换。')

    h(doc, '1.2 与概要设计的关系', level=2)
    para(doc, '概要设计回答"系统由哪些模块组成、它们大致如何协作"，详细设计回答"每个模块的关键类是什么、核心方法接受什么参数、返回什么、内部如何处理异常、如何与其它模块交换数据"。二者的章节编号遵循"概要第 X 章 → 详细第 Y 章"的对应关系：概要的第 4 章（功能模块设计）对应本文档第 3 章；概要的第 5 章（数据库设计）对应本文档第 4 章；概要的第 6 章（接口设计）对应本文档第 5 章；概要的第 7 章（关键流程）对应本文档第 8 章。')

    h(doc, '1.3 读者对象', level=2)
    bullet(doc, '后端开发：关注第 3、4、5、8、9、10 章，直接按类图与方法签名实现')
    bullet(doc, '前端开发：关注第 5、6、8 章，理解接口协议与 SSE 事件格式')
    bullet(doc, '测试人员：关注第 5、9、11 章，设计单元测试与端到端测试用例')
    bullet(doc, '维护与运维：关注第 9、10 章了解日志、并发限制、性能特性')

    h(doc, '1.4 术语约定', level=2)
    table(doc,
        ['术语', '含义'],
        [
            ['子任务 / item', 'batch_task_items 表中的一条记录，代表"(URL, POC)"或"(URL, 模板)"的最小执行单元'],
            ['流式接口 / SSE', '通过 text/event-stream 持续推送 JSON 事件的 HTTP 接口'],
            ['Provider', 'OOB 或空间测绘平台的供应商抽象，可在 interactsh / ceye 或 fofa / hunter / quake 间切换'],
            ['Runtime Params', '执行 url_with_params 模式 POC 时由用户补充的参数字典'],
            ['Helper', '平台在 POC 执行前注入的辅助模块（http_runtime、oob_runtime）'],
            ['笛卡尔积展开', '将 URL 列表与 POC 列表交叉生成所有 (URL, POC) 组合并写入 items 表的过程'],
        ],
        col_widths=[3, 12.5]
    )

    doc.add_page_break()

    # ========== 2 设计约束 ==========
    h(doc, '2 设计约束与全局依赖', level=1)

    h(doc, '2.1 设计约束', level=2)
    bullet(doc, '单进程部署：FastAPI + Uvicorn，默认监听 127.0.0.1:8000；批量任务通过线程池并发，不依赖消息队列与分布式调度')
    bullet(doc, 'SQLite 单写并发限制：所有写操作通过 contextmanager 提供的连接管理，事务提交后立即关闭，避免长事务锁')
    bullet(doc, '批量任务规模上限：MAX_URLS=200、MAX_POCS=50、MAX_TASK_ITEMS=2000、MAX_CONCURRENCY=5，超限在 create_task 阶段直接抛 ValueError')
    bullet(doc, 'LLM 调用超时与失败：不做自动重试，失败立即返回给前端，同时保留原始响应到 pocs/metadata/last_llm_response.json 便于排查')
    bullet(doc, 'POC 执行沙盒缺失：当前依赖 importlib + sys.modules 动态加载，未做沙箱隔离，执行前会做依赖静态检查，运行时打 requests 超时补丁')
    bullet(doc, '所有面向外部的 API Key 均保存在 pocs/*.json，文件权限由部署环境保证，前端仅展示脱敏预览')

    h(doc, '2.2 服务依赖拓扑', level=2)
    para(doc, '服务层内部存在以下调用关系，箭头方向表示"调用者 → 被调用者"：')
    code(doc, '''
routes.generate_poc
    ├─▶ llm_service.generate_initial_poc  ─▶ _call_llm_api_with_settings ─▶ AsyncOpenAI
    ├─▶ llm_service.review_generated_poc  ─▶ 同上（不同配置）
    ├─▶ poc_library_service.check_poc_dependencies ─▶ dependency_checker
    └─▶ poc_library_service.save_poc ─▶ SQLite + 文件系统

routes.execute_poc_by_id
    └─▶ poc_library_service.execute_poc
            ├─ python 分支 ─▶ _execute_python_poc
            │                   ├─▶ dependency_checker.check_python_file_dependencies
            │                   ├─▶ oob_service.create_client (注入 helper)
            │                   ├─▶ http_runtime.* (注入 helper)
            │                   └─▶ failure_classifier.classify_execution_outcome
            └─ nuclei 分支 ─▶ nuclei_service.scan_single

routes.create_batch_task
    └─▶ batch_task_service.create_task
            ├─▶ poc_library_service.get_poc_by_id (校验)
            └─▶ start_task ─▶ 新线程 ─▶ ThreadPoolExecutor
                                ├─▶ _execute_task_item ─▶ poc_library_service.execute_poc
                                └─▶ _execute_nuclei_task_item ─▶ nuclei_service.scan_single
'''.strip('\n'))

    h(doc, '2.3 关键全局单例', level=2)
    table(doc,
        ['变量名', '定义位置', '说明'],
        [
            ['llm_service', 'services/llm_service.py 模块末尾', 'LLMService 全局实例，进程启动时从 pocs/llm_config.json 加载配置'],
            ['poc_library_service', 'services/poc_library_service.py', 'PocLibraryService 全局实例，初始化时自动建表与迁移'],
            ['batch_task_service', 'services/batch_task_service.py', 'BatchTaskService 全局实例，维护 _worker_threads 与 _cancel_events'],
            ['nuclei_service', 'services/nuclei_service.py', 'NucleiService 全局实例，含模板缓存'],
            ['oob_service', 'services/oob_service.py', 'OOBService 全局实例，提供 provider_factories 注册点'],
            ['asset_source_service', 'services/asset_source_service.py', 'AssetSourceService 全局实例'],
        ],
        col_widths=[4, 5.5, 8]
    )

    doc.add_page_break()

    # ========== 3 业务服务层详细设计 ==========
    h(doc, '3 业务服务层详细设计', level=1)

    # ----- 3.1 LLMService -----
    h(doc, '3.1 LLMService —— 大模型生成服务', level=2)

    h(doc, '3.1.1 类属性', level=3)
    table(doc,
        ['属性', '类型', '默认值', '说明'],
        [
            ['config_file', 'Path', 'pocs/llm_config.json', '主模型配置文件路径'],
            ['review_config_file', 'Path', 'pocs/review_llm_config.json', '二次审核模型配置文件路径'],
            ['prompt_config_file', 'Path', 'prompts/poc_generation.yaml', 'Prompt 模板文件路径'],
            ['prompt_config', 'Dict', '加载自 YAML', 'Prompt 配置（含 user_prompt_template）'],
            ['api_key / model / api_base / temperature / max_tokens', 'str/str/str/float/int', '来自 DEFAULT_LLM_CONFIG 或 json 文件', '主模型运行参数'],
            ['client', 'AsyncOpenAI', '基于主模型配置初始化', 'LLM API 客户端'],
            ['review_api_key / review_model / review_api_base / review_temperature / review_max_tokens', '同上', '来自 review_llm_config.json', '二次审核模型运行参数'],
        ],
        col_widths=[5, 2.5, 4, 5]
    )

    h(doc, '3.1.2 公共方法签名', level=3)
    code(doc, '''
def update_config(self, api_key: str, model_id: str, base_url: str,
                  temperature: float = 0.7, max_tokens: Optional[int] = None) -> None
def get_current_config(self) -> Dict[str, str]
def update_review_config(self, api_key: str, model_id: str, base_url: str,
                         temperature: float = 0.3, max_tokens: Optional[int] = None) -> None
def get_current_review_config(self) -> Dict[str, str]

async def generate_initial_poc(
    self, vulnerability_info: str, target_info: Optional[str] = None
) -> Dict[str, Optional[str]]

async def review_generated_poc(
    self,
    vulnerability_info: str,
    target_info: Optional[str],
    initial_prompt: str,
    initial_result: Dict[str, Optional[str]],
) -> Dict[str, Optional[str]]
'''.strip('\n'))

    h(doc, '3.1.3 返回结构', level=3)
    para(doc, 'generate_initial_poc 与 review_generated_poc 均返回下述结构，可直接用于持久化或向前端响应：')
    table(doc,
        ['字段', '类型', '说明'],
        [
            ['verifiable', 'bool', 'True 表示可自动化；False 走 manual_guide'],
            ['vulnerability_name', 'str', '漏洞名称，含产品或 CVE'],
            ['vulnerability_type', 'str', 'sqli / xss / rce / ssrf 等'],
            ['original_vulnerability_info', 'str', '简化后的原始漏洞信息'],
            ['poc_code', 'str | None', '可验证时提供 scan 函数源码'],
            ['execution_mode', 'str', 'url_only / url_with_params / manual_guide'],
            ['verification_method', 'str', 'direct / oob / manual'],
            ['input_schema', 'List[Dict] | None', 'url_with_params 模式下必填'],
            ['manual_steps', 'Dict | None', 'manual_guide 模式下必填'],
            ['explanation', 'str', '代码说明或操作说明'],
        ],
        col_widths=[4, 2.5, 9]
    )

    h(doc, '3.1.4 关键算法：_parse_llm_json_response', level=3)
    para(doc, '本方法对 LLM 返回的 JSON 做健壮性解析，步骤如下：')
    code(doc, '''
1. 去除 Markdown 代码块标记（```json ... ``` 或 ``` ... ```）
2. 第一次尝试 json.loads(cleaned)
3. 若失败 → 调用 fix_control_chars_in_strings()：
     - 状态机扫描字符串，跟踪 in_string / escape_next
     - 字符串内部遇到真实 \\n \\r \\t 自动替换为转义形式
     - 字符串外部保留原样
4. 修复后再次 json.loads
5. 仍失败 → 抛出 Exception，并将原始内容保存到
   pocs/metadata/last_llm_response.json 供事后排查
6. 针对 poc_code 字段：去除 ```python 与 ``` 包裹；
   通过 _normalize_poc_code_string() 把字面量 "\\n" 反向替换为真实换行
   （仅当代码中没有真实换行且存在字面量转义时才处理，避免误伤）
'''.strip('\n'))

    h(doc, '3.1.5 关键算法：Prompt 二阶段组装', level=3)
    para(doc, '_build_prompt 优先使用 prompts/poc_generation.yaml 中的 user_prompt_template；如配置文件不存在则使用内置默认模板。模板采用字符串 replace 而非 str.format，是为了避免 Prompt 文本中其他 "{}" 被误解析。_build_review_prompt 则在原始信息基础上追加"第一次发送的完整 Prompt + 第一次的生成结果 JSON"，让审核模型在完整上下文下做改进。')

    # ----- 3.2 PocLibraryService -----
    h(doc, '3.2 PocLibraryService —— POC 库管理服务', level=2)

    h(doc, '3.2.1 主要方法签名', level=3)
    code(doc, '''
def save_poc(self,
             vuln_type: str, vuln_info: str,
             vuln_name: str = None, poc_code: str = None,
             explanation: str = "", poc_type: str = "python",
             tags: List[str] = None, metadata: dict = None,
             verifiable: bool = True,
             manual_steps: dict = None,
             execution_mode: Optional[str] = None,
             verification_method: Optional[str] = None,
             input_schema: Optional[List[Dict]] = None) -> int

def get_poc_by_id(self, poc_id: int) -> Optional[Dict]
def search_pocs(self, vuln_type=None, poc_type=None, keyword=None,
                verifiable=None, limit=50, offset=0) -> List[Dict]
def get_all_vuln_types(self) -> List[Dict]
def get_statistics(self) -> Dict
def execute_poc(self, poc_id: int, target_url: str,
                runtime_params: Optional[Dict[str, Any]] = None) -> Dict
def check_poc_dependencies(self, code: str) -> Dict[str, Any]
def delete_poc(self, poc_id: int) -> bool
'''.strip('\n'))

    h(doc, '3.2.2 关键算法：execute_poc 执行分派', level=3)
    code(doc, '''
输入: poc_id, target_url, runtime_params
1. 读取 poc_records 记录；找不到 → 返回 {"success": False, "error": "POC不存在"}
2. 若 verification_method == "oob"：
      runtime_status = oob_service.get_runtime_status()
      若 runtime_ready == False → 返回带 oob_error/not_configured 的结果
3. 更新 last_used
4. 根据 poc_type 分派：
   python  → _execute_python_poc(poc_file_path, target_url, runtime_params)
   nuclei  → 先检查 nuclei 可用性 → scan_single → _normalize_nuclei_result 统一结构
   其它    → 返回 {"success": False, "error": "不支持的POC类型"}
5. 无论成功失败，最终都通过 failure_classifier.classify_execution_outcome
   给结果附加 classification 字段
'''.strip('\n'))

    h(doc, '3.2.3 关键算法：_execute_python_poc 动态加载', level=3)
    code(doc, '''
1. _normalize_url(target_url) —— 统一补全 scheme、去除尾部 /
2. 检查 POC 文件存在性
3. 静态依赖预检：check_python_file_dependencies
   若 missing 不为空 → 返回 environment_error / missing_dependency
4. 动态构造模块：
   module_name = f"poc_module_{timestamp_ms}"   # 保证不重名
   spec = importlib.util.spec_from_file_location(module_name, poc_path)
   poc_module = importlib.util.module_from_spec(spec)
   sys.modules[module_name] = poc_module
5. 注入 helper 模块：
   sys.modules["oob_runtime"]  = <ModuleType 含 create_client/get_oob_client>
   sys.modules["http_runtime"] = <ModuleType 含 HTTPRuntimeClient 等>
   同时把这些符号直接绑定到 poc_module 命名空间，兼容老式 POC
   把 runtime_params / runtime_input / input_params / get_runtime_param 注入模块
6. spec.loader.exec_module(poc_module)
7. 获取 scan 函数并执行：
   - 打上 requests 超时补丁（_patch_requests_timeout）
   - _invoke_scan(scan, normalized_url, runtime_params)：
     先尝试 scan(url, **runtime_params)；
     若 TypeError 提示参数不匹配 → 回退到 scan(url)
8. finally 清理 sys.modules 中刚注入的三个模块名，避免长期持有
9. 若返回值非 dict → 规范化为 {"vulnerable": False, "reason": "格式不正确", ...}
10. 若识别到网络失败特征 → 转为 {"success": False, ...} 以便前端归类
'''.strip('\n'))

    h(doc, '3.2.4 关键算法：POC 文件命名与落盘', level=3)
    para(doc, 'save_poc 使用 _normalize_vuln_name 生成人类可读名、_build_file_stem 生成安全文件名，最终写入：')
    code(doc, '''
文件名规则:
  <safe_vuln_name>_<YYYYMMDD_HHMMSS>_<sha1_prefix6>.py
示例:
  SQL注入_20260418_213000_27c38c.py

落盘位置:
  pocs/python/<filename>.py      ← 实际可执行脚本
  pocs/metadata/<filename>.json  ← LLM 原始生成快照（含 execution_mode、input_schema 等）

数据库:
  INSERT INTO poc_records (vuln_type, vuln_name, vuln_description,
                           poc_type, poc_file_path, tags, metadata,
                           verifiable, manual_steps,
                           execution_mode, verification_method, input_schema)
  VALUES (...)
返回自增 id
'''.strip('\n'))

    # ----- 3.3 BatchTaskService -----
    h(doc, '3.3 BatchTaskService —— 批量任务编排服务', level=2)

    h(doc, '3.3.1 类常量与可调参数', level=3)
    table(doc,
        ['常量', '值', '说明'],
        [
            ['MAX_URLS', '200', '单任务 URL 上限'],
            ['MAX_POCS', '50', '单任务 POC / 模板上限'],
            ['MAX_TASK_ITEMS', '2000', '笛卡尔积后子任务上限'],
            ['DEFAULT_CONCURRENCY', '3', '未指定时使用的并发数'],
            ['MAX_CONCURRENCY', '5', '用户指定并发数的上限'],
        ],
        col_widths=[4, 2, 9.5]
    )

    h(doc, '3.3.2 主要方法签名', level=3)
    code(doc, '''
def create_task(self, target_urls: List[str], poc_ids: List[int],
                concurrency: Optional[int] = None) -> Dict
def create_nuclei_task(self, target_urls: List[str],
                       template_paths: List[str],
                       concurrency: Optional[int] = None) -> Dict
def start_task(self, task_id: int) -> None       # 启动后台线程
def cancel_task(self, task_id: int) -> bool
def list_tasks(self, limit=20, offset=0, status=None,
               result_filter=None, keyword=None,
               sort_by="created_at", sort_order="desc") -> Dict
def get_task(self, task_id: int) -> Optional[Dict]
def get_task_items(self, task_id, status=None, keyword=None,
                   limit=200, offset=0) -> Dict
def get_task_item_detail(self, task_id: int, item_id: int) -> Optional[Dict]
def export_task_report(self, task_id: int, report_format: str) -> Tuple[str, str, bytes]
'''.strip('\n'))

    h(doc, '3.3.3 状态机设计', level=3)
    para(doc, '任务及子任务生命周期使用状态字段管理，状态迁移如下：')
    code(doc, '''
batch_tasks.status:
  pending ──start_task──▶ running ──全部完成──▶ completed
                                └──cancel──▶ cancelled
                                └──异常──▶ failed

batch_task_items.status:
  pending ──分派执行──▶ running ──outcome.success == True──▶ success
                              └──outcome.success == False──▶ failed
                              └──task cancelled──▶ cancelled
'''.strip('\n'))

    h(doc, '3.3.4 关键算法：_run_task 调度循环', level=3)
    code(doc, '''
def _run_task(self, task_id, cancel_event):
    1. 读取任务与全部 pending 子任务
    2. 将 batch_tasks.status 从 pending 改为 running，
       started_at = COALESCE(started_at, CURRENT_TIMESTAMP)
    3. 创建 ThreadPoolExecutor(max_workers=concurrency)
    4. 主循环：
       while True:
         if cancel_event.is_set(): break
         # 保持队列填满
         while len(futures) < concurrency:
             item = next(pending_iter)  # StopIteration 退出内层循环
             _mark_item_running(item.id)
             futures[executor.submit(_execute_task_item, item)] = item.id
         if not futures: break
         done, _ = wait(futures, FIRST_COMPLETED, timeout=0.5)
         for future in done:
             outcome = future.result() 或 统一异常包装
             _store_item_result(item_id, outcome)
             _refresh_task_stats(task_id)
    5. 取消场景：剩余 futures 统一写 cancelled
    6. _finalize_task(task_id, cancelled)：写 completed/cancelled/failed 与 finished_at
    7. finally: 从 _worker_threads 与 _cancel_events 中移除
'''.strip('\n'))

    h(doc, '3.3.5 关键算法：统计刷新与报告导出', level=3)
    code(doc, '''
_refresh_task_stats(task_id):
  SELECT
      COUNT(*) AS total,
      SUM(CASE status WHEN 'success' THEN 1 ELSE 0 END) AS success_items,
      SUM(CASE status WHEN 'failed'  THEN 1 ELSE 0 END) AS failed_items,
      SUM(CASE WHEN json_extract(result_json,'$.vulnerable') = 1 THEN 1 ELSE 0 END)
                                                      AS vulnerable_items,
      SUM(CASE status IN ('success','failed','cancelled') THEN 1 ELSE 0 END)
                                                      AS completed_items
  FROM batch_task_items WHERE task_id = ?
  UPDATE batch_tasks SET completed_items=?, success_items=?,
                         failed_items=?, vulnerable_items=?
  WHERE id=?

export_task_report(task_id, format):
  payload = build_task_report_payload(task_id)   # 包含任务概览 + 子任务明细
  if format == "json": dump 为 JSON bytes
  if format == "html": _render_html_report 生成带样式的 HTML
  if format == "txt":  _render_text_report 生成纯文本
  返回 (filename, content_type, content_bytes)
'''.strip('\n'))

    # ----- 3.4 NucleiService -----
    h(doc, '3.4 NucleiService —— Nuclei 兼容扫描服务', level=2)

    h(doc, '3.4.1 主要方法签名', level=3)
    code(doc, '''
def check_nuclei_available(self) -> Dict       # 子进程执行 nuclei -version
def get_folder_structure(self) -> List[Dict]
def get_templates_by_folder(self, folder="", page=1, page_size=100,
                            keyword="") -> Tuple[List[Dict], int]
def get_template_content(self, template_path: str) -> Optional[str]
def get_total_template_count(self) -> int
def scan_single(self, target_url: str, template_path: str,
                timeout: int = 60) -> Dict
def scan_multiple(self, target_url: str, template_paths: List[str],
                  timeout: int = 120) -> Dict
def scan_folder(self, target_url: str, folder: str,
                timeout: int = 300) -> Dict
async def scan_stream_async(self, target_url, template_paths=None,
                            folder=None, timeout=120)  # SSE 异步生成器
def clear_cache(self) -> None
'''.strip('\n'))

    h(doc, '3.4.2 关键算法：_parse_template_fast', level=3)
    para(doc, '为避免在大模板库（数万文件）下完整 YAML 解析导致慢启动，nuclei_service 使用轻量解析策略：')
    code(doc, '''
打开模板 YAML，逐行扫描：
  - 遇到顶层键 id: / info: / name: / author: / severity: / tags: 时
    只截取第一个冒号右侧的值，strip 引号
  - info 块内采用缩进识别，遇到空行或下一个顶层键即停止
  - description 仅记录首行，后续多行描述忽略
  - tags 支持逗号分隔字符串，按 "," 分割
结果缓存到内存字典，key = 相对路径
clear_cache() 清空后会在下次 get_templates_by_folder 时重建
'''.strip('\n'))

    h(doc, '3.4.3 关键算法：scan_stream_async 子进程流式解析', level=3)
    code(doc, '''
使用 asyncio.create_subprocess_exec 启动:
  nuclei -u <target> -t <template>... -jsonl -silent -nc ...

异步读取 stdout 行：
  for line in iter:
      try:
          record = json.loads(line)
          yield {"type": "finding", "data": _format_finding(record)}
      except JSONDecodeError:
          yield {"type": "log", "message": line}
stderr 汇总到 errors
进程结束后 yield {"type": "done", "summary": {...}}

超时时 kill 子进程并 yield {"type": "error", "message": "timeout"}
'''.strip('\n'))

    # ----- 3.5 OOB -----
    h(doc, '3.5 OOBService 与 OOB 客户端', level=2)

    h(doc, '3.5.1 BaseOOBClient 抽象接口', level=3)
    code(doc, '''
class BaseOOBClient:
    def build_probe(self, protocol: str = "http",
                    length: int = 10, value: str = "") -> Dict[str, str]:
        \"\"\"返回 {"url": "<回连地址>", "flag": "<一次性标识>"}\"\"\"
        raise NotImplementedError

    def verify(self, flag: str,
               protocol: str = "http") -> Dict[str, object]:
        \"\"\"返回 {"hit": bool, "records": [...], "raw": ...}\"\"\"
        raise NotImplementedError
'''.strip('\n'))

    h(doc, '3.5.2 InteractshClient 实现要点', level=3)
    bullet(doc, '注册阶段：客户端生成 RSA 2048 密钥对；将 correlation_id + 公钥上报 interactsh server；server 回传由公钥加密的 AES 密钥与 secret_key')
    bullet(doc, 'build_probe：拼接 <random_sub>.<correlation_id>.<server_domain>，并返回 flag（用于后续 verify 关联）')
    bullet(doc, 'verify：带 secret_key 轮询 /poll 接口，收到密文数组后用本地 RSA 私钥解出 AES 密钥，再 AES-CFB 解密每条记录；按 protocol 过滤 http/dns')
    bullet(doc, '轮询策略：poll_interval 秒一次，最多 max_polls 次；任意一次命中即返回')

    h(doc, '3.5.3 CEyeClient 实现要点', level=3)
    bullet(doc, 'build_probe：identifier = <随机前缀>.<token>.ceye.io；返回对应 url / flag')
    bullet(doc, 'verify：调用 ceye API ({base_url}/record/{flag}?type=http|dns&token=xxx) 拉取记录；按 poll_interval 轮询，max_polls 次退出')
    bullet(doc, 'CEye 的 HTTP 回显延时约 1-2s，建议 poll_interval ≥ 1.0、max_polls ≥ 3')

    h(doc, '3.5.4 OOBService 对外方法', level=3)
    code(doc, '''
def update_config(self, enabled: bool, provider: str, ...)  # provider: interactsh/ceye
def get_current_config(self) -> Dict
def get_runtime_status(self) -> Dict
    # 返回 {dependency_ready, dependency_error, runtime_ready, runtime_error}
def create_client(self) -> BaseOOBClient
    # POC 内通过 get_oob_client() / create_oob_client() 调用
'''.strip('\n'))

    # ----- 3.6 FailureClassifier -----
    h(doc, '3.6 FailureClassifier —— 失败归因', level=2)
    para(doc, 'classify_execution_outcome 接收统一 outcome 字典，按优先级匹配关键字并返回 {failure_category, failure_code, failure_stage, retryable}：')
    table(doc,
        ['优先级', '类别', '失败码', '触发关键字（大小写不敏感）', '可重试'],
        [
            ['1', '—（成功）', '—', 'success=True 且 vulnerable=True', '否'],
            ['2', 'oob_error', 'not_configured', '"oob 运行环境不可用" + "未启用/not configured"', '否'],
            ['2', 'oob_error', 'runtime_unavailable', 'interactsh / ceye 关键字', '否'],
            ['3', 'nuclei_error', 'engine_unavailable', '"nuclei" + 不可用/未安装', '否'],
            ['3', 'nuclei_error', 'scan_timeout', '"nuclei" + 超时/timeout', '是'],
            ['3', 'nuclei_error', 'scan_failed', '其它 nuclei 失败', '否'],
            ['4', 'code_error', 'syntax_error', 'syntaxerror / invalid syntax', '否'],
            ['5', 'environment_error', 'missing_dependency', '"缺少依赖" / missing_dependency', '否'],
            ['5', 'environment_error', 'import_error', 'modulenotfounderror / no module named', '否'],
            ['6', 'network_error', 'dns / connect_timeout / read_timeout / connection_refused / connection_reset / connect_failed / timeout', 'NETWORK_CODES 元组', '是'],
            ['7', 'unknown', 'unclassified', 'success=False 且无上述命中', '否'],
            ['8', 'not_vulnerable', 'no_evidence_found', 'success=True 且 vulnerable=False', '否'],
        ],
        col_widths=[1.5, 2.5, 3, 8, 1.5]
    )

    # ----- 3.7 DependencyChecker -----
    h(doc, '3.7 DependencyChecker —— POC 依赖静态预检', level=2)
    h(doc, '3.7.1 对外接口', level=3)
    code(doc, '''
def check_python_code_dependencies(code: str) -> Dict[str, Any]
def check_python_file_dependencies(file_path: str | Path) -> Dict[str, Any]
返回结构:
  {"ok": bool, "imports": [...], "missing": [...], "summary": str,
   "parse_error": 可选}
''')

    h(doc, '3.7.2 算法流程', level=3)
    code(doc, '''
1. ast.parse(code) → 抽象语法树；解析失败时返回 ok=False 并附带 parse_error
2. 遍历 AST，收集根模块名：
   - ast.Import 节点 → alias.name.split(".")[0]
   - ast.ImportFrom 节点 → node.module.split(".")[0]；
     当 node.level > 0 且 module 为空时跳过（相对 import）
3. 对每个根模块 _is_module_available(name)：
   - 命中 PLATFORM_PROVIDED_MODULES = {"http_runtime", "oob_runtime"} → True
   - sys.stdlib_module_names 中 → True
   - importlib.util.find_spec(name) is not None → True
   - 否则 False
4. missing = [未命中的模块]
5. 返回 {ok: not missing, imports, missing, summary}
'''.strip('\n'))

    # ----- 3.8 HTTP Runtime -----
    h(doc, '3.8 HTTP Runtime —— 统一请求运行时', level=2)
    h(doc, '3.8.1 HTTPRuntimeClient', level=3)
    code(doc, '''
@dataclass
class HTTPRuntimeClient:
    timeout: int = 6
    verify: bool = False
    allow_redirects: bool = True
    headers: Dict[str, str] = {"User-Agent": "AI-POC/1.0", ...}
    session: requests.Session

    def request(method, url, **kwargs)
    def get/post/put/delete(url, **kwargs)
    def raw_request(raw, use_ssl=False, timeout=None)  # → send_raw_http
'''.strip('\n'))

    h(doc, '3.8.2 send_raw_http 算法', level=3)
    code(doc, '''
1. 切分 raw，第一行拆为 method / path / version
2. 逐行扫描 headers，直到空行；剩余为 body
3. 取 Host 头，解析 hostname:port；端口缺省按 use_ssl 决定
4. 构造 request_text，将 \\n 替换为 \\r\\n，若无空行结尾则补 \\r\\n\\r\\n
5. socket.create_connection((hostname, port), timeout)
6. 若 use_ssl：SSLContext.wrap_socket（不校验证书）
7. sendall request_text；while recv(4096) → 拼接
8. decode 优先 utf-8，失败回退 latin1 replace
9. 返回 {method, host, port, path, raw_response, status_line, body}
'''.strip('\n'))

    h(doc, '3.8.3 _patch_requests_timeout', level=3)
    code(doc, '''
目的: 为运行期的 requests.Session.request 注入默认超时，避免 POC 卡死

实现:
  original = requests.Session.request
  def request_with_timeout(self, method, url, **kwargs):
      kwargs.setdefault("timeout", 8)
      return original(self, method, url, **kwargs)
  requests.Session.request = request_with_timeout
  return restore   # 执行完 scan 后调用以还原
'''.strip('\n'))

    # ----- 3.9 资产导入 -----
    h(doc, '3.9 AssetSourceService —— 空间测绘导入', level=2)
    table(doc,
        ['Provider', '查询接口', '关键参数', '标准化输出'],
        [
            ['fofa', 'GET https://fofa.info/api/v1/search/all', 'email、key、qbase64(query)、page', 'host → http(s)://host'],
            ['hunter', 'GET https://hunter.qianxin.com/openApi/search', 'api-key、search(base64)、page', 'url 字段直接使用'],
            ['quake', 'POST https://quake.360.cn/api/v3/search/quake_service', 'token、query、start、size', '拼接 ip:port → http://ip:port'],
        ],
        col_widths=[2, 6, 4, 3.5]
    )
    para(doc, '服务内部按 provider 分派对应查询函数，统一返回 {success, provider, query, total, urls: [...]}。所有 HTTP 请求带 30s 超时，失败时抛 ValueError 以便上层返回 400 友好错误。')

    doc.add_page_break()

    # ========== 4 数据库详细设计 ==========
    h(doc, '4 数据模型与数据库详细设计', level=1)

    h(doc, '4.1 DDL 全量 SQL', level=2)
    code(doc, '''
CREATE TABLE IF NOT EXISTS poc_records (
    id                  INTEGER PRIMARY KEY AUTOINCREMENT,
    vuln_type           TEXT    NOT NULL,
    vuln_name           TEXT    NOT NULL,
    vuln_description    TEXT,
    poc_type            TEXT    DEFAULT 'python',
    poc_file_path       TEXT    NOT NULL,
    create_time         TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    last_used           TIMESTAMP,
    tags                TEXT,
    metadata            TEXT,
    verifiable          BOOLEAN DEFAULT 1,
    manual_steps        TEXT,
    execution_mode      TEXT,          -- 字段迁移时按需补齐
    verification_method TEXT,
    input_schema        TEXT
);
CREATE INDEX IF NOT EXISTS idx_vuln_type   ON poc_records(vuln_type);
CREATE INDEX IF NOT EXISTS idx_poc_type    ON poc_records(poc_type);
CREATE INDEX IF NOT EXISTS idx_create_time ON poc_records(create_time DESC);
CREATE INDEX IF NOT EXISTS idx_verifiable  ON poc_records(verifiable);

CREATE TABLE IF NOT EXISTS batch_tasks (
    id                INTEGER PRIMARY KEY AUTOINCREMENT,
    task_type         TEXT    NOT NULL DEFAULT 'poc_batch',
    mode              TEXT    NOT NULL,
    status            TEXT    NOT NULL DEFAULT 'pending',
    total_items       INTEGER NOT NULL DEFAULT 0,
    completed_items   INTEGER NOT NULL DEFAULT 0,
    success_items     INTEGER NOT NULL DEFAULT 0,
    failed_items      INTEGER NOT NULL DEFAULT 0,
    vulnerable_items  INTEGER NOT NULL DEFAULT 0,
    concurrency       INTEGER NOT NULL DEFAULT 3,
    config_json       TEXT,
    error             TEXT,
    created_at        TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    started_at        TIMESTAMP,
    finished_at       TIMESTAMP
);
CREATE INDEX IF NOT EXISTS idx_batch_tasks_status ON batch_tasks(status);

CREATE TABLE IF NOT EXISTS batch_task_items (
    id                 INTEGER PRIMARY KEY AUTOINCREMENT,
    task_id            INTEGER NOT NULL,
    poc_id             INTEGER NOT NULL,
    target_url         TEXT    NOT NULL,
    engine_type        TEXT    NOT NULL DEFAULT 'poc',
    template_path      TEXT,
    status             TEXT    NOT NULL DEFAULT 'pending',
    result_json        TEXT,
    error              TEXT,
    failure_category   TEXT,
    failure_code       TEXT,
    failure_stage      TEXT,
    retryable          INTEGER NOT NULL DEFAULT 0,
    created_at         TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    started_at         TIMESTAMP,
    finished_at        TIMESTAMP,
    FOREIGN KEY(task_id) REFERENCES batch_tasks(id)
);
CREATE INDEX IF NOT EXISTS idx_batch_task_items_task_id ON batch_task_items(task_id);
CREATE INDEX IF NOT EXISTS idx_batch_task_items_status  ON batch_task_items(status);
'''.strip('\n'))

    h(doc, '4.2 字段兼容迁移策略', level=2)
    para(doc, '由于项目经历了多次字段演进（verifiable、execution_mode、input_schema 等陆续加入），init_database 与 _ensure_batch_task_item_columns 采用如下迁移流程：')
    code(doc, '''
迁移伪代码:
  PRAGMA table_info(<table>);
  existing_cols = {row[1] for row in fetchall()}
  for col, ddl in REQUIRED_COLUMNS.items():
      if col not in existing_cols:
          ALTER TABLE <table> ADD COLUMN <col> <ddl>;
  (不执行 DROP / RENAME，保证老库可直接升级)
'''.strip('\n'))

    h(doc, '4.3 连接管理与并发', level=2)
    code(doc, '''
@contextmanager
def get_db_connection(self):
    conn = sqlite3.connect(self.db_path, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()

要点:
  - check_same_thread=False  —— 线程池中可复用
  - row_factory=sqlite3.Row —— 取值时用字典访问
  - 显式 commit/rollback，无隐式事务
  - SQLite 单文件写锁限制：所有写路径都通过 contextmanager 快进快出
'''.strip('\n'))

    h(doc, '4.4 POC 文件与元数据布局', level=2)
    table(doc,
        ['路径', '内容', '写入时机'],
        [
            ['pocs/python/<safe_name>.py', 'Python 可执行 POC 脚本', 'save_poc()'],
            ['pocs/metadata/<safe_name>.json', 'LLM 原始生成结果快照', 'save_poc()'],
            ['pocs/metadata/last_llm_response.json', '最近一次主模型原始返回', '_call_llm_api_with_settings 总是覆盖'],
            ['pocs/metadata/last_llm_review_response.json', '最近一次审核模型原始返回', '同上'],
            ['pocs/batch_results/<item_id>.json', '批量子任务详细结果落盘', '_write_detail_file()'],
            ['pocs/nuclei/templates/...', 'Nuclei 官方模板库', '用户手动维护或解压分发'],
            ['pocs/llm_config.json', '主模型配置', 'update_config()'],
            ['pocs/review_llm_config.json', '审核模型配置', 'update_review_config()'],
            ['pocs/oob_config.json', 'OOB 配置', 'OOBService.update_config()'],
            ['pocs/asset_source_config.json', '空间测绘配置', 'AssetSourceService.update_provider_config()'],
        ],
        col_widths=[5, 5, 5.5]
    )

    doc.add_page_break()

    # ========== 5 接口详细设计 ==========
    h(doc, '5 接口详细设计', level=1)

    h(doc, '5.1 通用约定', level=2)
    bullet(doc, '前缀统一为 /api；静态资源在 / /css /js /assets')
    bullet(doc, 'Content-Type 默认 application/json；POC 生成与 Nuclei 流式扫描使用 text/event-stream；报告下载使用 Content-Disposition: attachment')
    bullet(doc, '成功响应：业务层在 body 中返回 success:true + 其它字段；失败由 FastAPI 统一包装为 {"detail": "..."} 并带有对应 HTTP 状态码')
    bullet(doc, '路径参数 {poc_id}、{task_id} 均为正整数；非法值由 FastAPI 自动返回 422')

    h(doc, '5.2 POST /api/generate-poc', level=2)
    table(doc,
        ['字段', '位置', '类型', '必填', '说明'],
        [
            ['vulnerability_info', 'body', 'str', '是', '漏洞信息/CVE/HTTP 报文'],
            ['target_info', 'body', 'str', '否', '目标系统附加说明（已弱化）'],
            ['enable_second_review', 'body', 'bool', '否，默认 false', '是否调用二次审核模型'],
        ],
        col_widths=[5, 2, 2.5, 2, 5]
    )
    para(doc, '响应：text/event-stream，按顺序推送以下事件，最后以 "data: [DONE]" 结束。')
    table(doc,
        ['事件 type', '阶段 step', '负载字段', '示意'],
        [
            ['status', '0', 'message', '"开始生成流程..."'],
            ['status', '1', 'message/status=active|completed', '"正在使用 {model} 生成POC代码..."'],
            ['status', '2', 'message/status', '"正在使用 {review_model} 进行二次审核..." 或 "未启用，已跳过"'],
            ['result', '—', 'data: PocResponse', '完整 PocResponse 模型（见模型 schemas）'],
            ['error', '—', 'data: PocResponse(success=false)', '生成失败时替代 result'],
        ],
        col_widths=[2, 1.5, 4.5, 7.5]
    )

    h(doc, '5.3 POC 库接口组', level=2)
    table(doc,
        ['方法 路径', '请求参数', '响应关键字段', '错误码'],
        [
            ['GET /api/pocs/search', 'vuln_type, poc_type, keyword, verifiable, limit, offset', 'success, total, pocs:[PocRecord]', '500'],
            ['GET /api/pocs/statistics', '—', 'success, statistics:{total, by_vuln_type, by_poc_type, by_verifiable}', '500'],
            ['GET /api/pocs/vuln-types', '—', 'success, vuln_types:[{type, count}]', '500'],
            ['GET /api/pocs/{poc_id}', '—', 'success, poc', '404 POC 不存在'],
            ['GET /api/pocs/{poc_id}/code', '—', 'success, code (字符串)', '404 POC/文件不存在 ｜ 500 读取失败'],
            ['POST /api/pocs/{poc_id}/execute', 'target_url, runtime_params', 'success, poc_id, target_url, result, classification', '404 POC 不存在 ｜ 500 执行异常'],
            ['DELETE /api/pocs/{poc_id}', '—', 'success, message', '404 POC 不存在'],
        ],
        col_widths=[4.5, 4, 4.5, 2.5]
    )

    h(doc, '5.4 批量任务接口组', level=2)
    table(doc,
        ['方法 路径', '请求参数', '响应关键字段', '错误码'],
        [
            ['POST /api/batch-tasks', 'target_urls:[str], poc_ids:[int], concurrency', 'success, task', '400 参数非法/超限 ｜ 500 其它'],
            ['GET /api/batch-tasks', 'limit, offset, status, result_filter(hit/failed/clean), keyword, sort_by, sort_order', 'success, total, tasks', '500'],
            ['GET /api/batch-tasks/{task_id}', '—', 'success, task', '404'],
            ['GET /api/batch-tasks/{task_id}/items', 'status, keyword, limit, offset', 'success, total, items', '404 / 500'],
            ['GET /api/batch-tasks/{task_id}/items/{item_id}/detail', '—', 'success, item (含完整 result_json)', '404 / 500'],
            ['POST /api/batch-tasks/{task_id}/cancel', '—', 'success, message, task', '404'],
            ['GET /api/batch-tasks/{task_id}/export', 'format=html|json|txt', '文件下载（Content-Disposition: attachment）', '400 非法 format ｜ 404 / 500'],
        ],
        col_widths=[4.5, 4, 4.5, 2.5]
    )

    h(doc, '5.5 Nuclei 接口组', level=2)
    table(doc,
        ['方法 路径', '请求参数', '响应关键字段', '错误码'],
        [
            ['GET /api/nuclei/status', '—', 'available, version, path, error, templates_count', '—'],
            ['GET /api/nuclei/folders', '—', 'success, folders:[{name, count}]', '500'],
            ['GET /api/nuclei/templates', 'folder, page, page_size, keyword', 'success, templates, total, page, page_size, total_pages', '500'],
            ['GET /api/nuclei/template/content', 'path', 'success, content', '404'],
            ['POST /api/nuclei/scan', 'target_url, template_paths/folder, timeout', 'NucleiScanResponse', '500'],
            ['POST /api/nuclei/scan/stream', '同上', 'SSE 事件：progress / finding / log / done / error', '—'],
            ['POST /api/nuclei/tasks', 'target_urls, template_paths/folder, concurrency', 'success, task（进入 batch_tasks 体系）', '400 / 500'],
            ['POST /api/nuclei/cache/clear', '—', 'success, message', '500'],
        ],
        col_widths=[4.5, 4, 4.5, 2.5]
    )

    h(doc, '5.6 配置接口组', level=2)
    table(doc,
        ['方法 路径', '行为', '说明'],
        [
            ['GET/POST /api/config/llm', '查询 / 更新主生成模型配置', '更新成功后立即热重建 AsyncOpenAI 客户端'],
            ['GET/POST /api/config/review-llm', '查询 / 更新审核模型配置', '同上但用于 review_*'],
            ['GET/POST /api/config/oob', '查询 / 更新 OOB 配置', '更新后调用 _save_config_to_file 持久化'],
            ['GET/POST /api/config/asset-sources', '查询 / 更新空间测绘凭证', 'provider 必须在 fofa/hunter/quake 中'],
            ['POST /api/asset-sources/import', '从测绘平台导入目标', '返回 urls 列表，前端直接灌入批量任务'],
        ],
        col_widths=[5, 4, 6.5]
    )

    h(doc, '5.7 错误码对照表', level=2)
    table(doc,
        ['HTTP 状态', '使用场景', '响应体示例'],
        [
            ['200', '业务成功', '{"success": true, ...}'],
            ['200', '业务失败但接口正常（如 POC 执行失败）', '{"success": false, "error": "目标不可达"}'],
            ['400', '入参校验失败（create_task 超限、非法 format）', '{"detail": "URL数量超出限制..."}'],
            ['404', '资源不存在（POC/任务/模板/子任务）', '{"detail": "POC不存在"}'],
            ['422', 'FastAPI 自动的 Pydantic 校验失败', '{"detail": [{"loc":[...], "msg":"..."}, ...]}'],
            ['500', '内部异常，已被 try/except 捕获', '{"detail": "...原始异常信息..."}'],
        ],
        col_widths=[2, 6.5, 7]
    )

    doc.add_page_break()

    # ========== 6 前端详细设计 ==========
    h(doc, '6 前端详细设计', level=1)

    h(doc, '6.1 页面结构', level=2)
    para(doc, '系统提供两套前端入口。主入口 index.html 采用单页应用（SPA）模式，通过 Tab 切换"POC 生成 / POC 库 / 检测记录 / 系统配置"四个面板；Nuclei 扫描器则独立页面 nuclei.html，以便保留与主界面不同的布局与状态。')

    h(doc, '6.2 JS 模块职责', level=2)
    table(doc,
        ['文件', '核心函数/入口', '说明'],
        [
            ['js/app.js', 'initGeneratePanel() / openGeneratePoc() / listenGenerateStream()', '生成页：表单收集 → fetch POST → EventSource 读流 → 实时 UI 更新'],
            ['js/app.js', 'openLlmConfigModal() / saveLlmConfig()', '系统配置弹窗（LLM/审核/OOB/资产）'],
            ['js/library.js', 'initLibraryPanel() / loadPocList() / renderPocCards()', 'POC 库：搜索筛选、分页、卡片渲染'],
            ['js/library.js', 'openExecuteDialog() / submitExecution()', '单条执行：按 execution_mode 决定是否弹参数表单'],
            ['js/library.js', 'createBatchTask() / pollBatchTask() / exportReport()', '批量任务：提交 → 2s 轮询进度 → 下载报告'],
            ['js/nuclei.js', 'loadFolders() / loadTemplates() / selectTemplate() / scanStream()', 'Nuclei：文件夹树 → 模板列表 → SSE 扫描'],
        ],
        col_widths=[3, 5.5, 7]
    )

    h(doc, '6.3 与后端交互模式', level=2)
    table(doc,
        ['场景', '机制', '实现细节'],
        [
            ['POC 生成', 'SSE (EventSource)', '接收 status/result/error 事件，按 type 分发到 UI；超时/断连时提示用户重试'],
            ['Nuclei 流扫描', 'SSE', '同上；finding 事件追加到结果面板'],
            ['批量任务进度', 'HTTP 轮询', '提交后 2 秒间隔请求 /api/batch-tasks/{id}，status ∈ {completed, cancelled, failed} 时停止；UI 更新进度条与统计'],
            ['POC 库搜索', 'HTTP GET', '输入防抖 300ms；limit/offset 分页；结果为空给出引导提示'],
            ['报告下载', 'window.open / a[download]', '后端返回 Content-Disposition，浏览器自动下载'],
        ],
        col_widths=[3.5, 3, 9]
    )

    h(doc, '6.4 关键交互序列', level=2)
    code(doc, '''
[生成页提交]
user  ─ click "生成" ──▶  app.js:openGeneratePoc
                              │ POST /api/generate-poc (fetch, keepalive)
                              ▼
                         后端 SSE 事件流
                              │ status/step=1
                              │ status/step=2
                              │ result  ── 渲染 POC 代码 + 标签
                              │ [DONE]
                              ▼
                         显示"已保存到 POC 库"并提示跳转
'''.strip('\n'))

    code(doc, '''
[批量任务轮询]
library.js:createBatchTask
  └─ POST /api/batch-tasks
        ├─ 成功 → 进入进度面板
        │      └─ setInterval 2000ms: GET /api/batch-tasks/{id}
        │             ├─ running  → 刷新进度条/统计
        │             ├─ completed/cancelled/failed → clearInterval, 弹出导出按钮
        │             └─ 任意失败 → 显示错误且保留轮询（用户可关闭）
        └─ 失败 → 保留表单状态，弹出错误提示
'''.strip('\n'))

    doc.add_page_break()

    # ========== 7 POC 脚本规范 ==========
    h(doc, '7 POC 脚本规范', level=1)

    h(doc, '7.1 函数签名', level=2)
    code(doc, '''
def scan(url, **kwargs):
    \"\"\"
    - url: 已由 _normalize_url 标准化的 target URL
    - kwargs: url_with_params 模式下注入的 runtime_params；url_only 时为 {}
    返回值必须是 dict，至少包含以下字段：
        vulnerable: bool
        reason:     str
        details:    任意可序列化对象
    \"\"\"
'''.strip('\n'))
    para(doc, '执行引擎会优先尝试 scan(url, **kwargs)；若函数签名不支持 kwargs（老 POC 只有 def scan(url):），则捕获 TypeError 并回退到 scan(url)。')

    h(doc, '7.2 平台注入的 Helper', level=2)
    table(doc,
        ['注入名', '类型', '用途'],
        [
            ['create_http_client() / get_http_client()', 'HTTPRuntimeClient', '统一 HTTP 客户端，默认 6s 超时、跟随跳转'],
            ['http_request(method, url, **kw)', 'function', '一次性请求'],
            ['send_raw_http(raw, use_ssl, timeout)', 'function', '原始 HTTP 报文发送'],
            ['get_oob_client() / create_oob_client()', 'BaseOOBClient', 'OOB 探针构造与 verify'],
            ['get_runtime_param(name, default=None)', 'function', '读取 runtime_params 的便捷方法'],
            ['runtime_params / runtime_input / input_params', 'dict', '用户填写的参数字典（三者别名）'],
        ],
        col_widths=[5.5, 3, 7]
    )

    h(doc, '7.3 input_schema 规范（url_with_params 模式）', level=2)
    code(doc, '''
[
  {
    "name": "cookie",
    "label": "登录 Cookie",
    "type": "textarea",        // text/password/textarea/json/select/checkbox
    "required": true,
    "description": "请粘贴完整的 Cookie 头",
    "placeholder": "PHPSESSID=xxx; ...",
    "options": [ {"value": "admin", "label": "管理员"} ]  // select 时提供
  }
]
'''.strip('\n'))

    h(doc, '7.4 manual_steps 规范（manual_guide 模式）', level=2)
    code(doc, '''
{
  "required_tools": ["Burp Suite", "浏览器"],
  "steps": [
    {"index": 1, "title": "登录后台", "description": "...", "expected": "..."},
    ...
  ],
  "verification": {
    "criteria": "响应包含关键字 XXX",
    "screenshot_hint": "建议截图保存"
  }
}
'''.strip('\n'))

    doc.add_page_break()

    # ========== 8 关键流程时序设计 ==========
    h(doc, '8 关键流程时序设计', level=1)

    h(doc, '8.1 POC 生成（含二次审核）', level=2)
    code(doc, '''
Frontend        FastAPI            LLMService        PocLibService     LLM
   │ POST generate-poc │                 │                  │             │
   │──────────────────▶│                 │                  │             │
   │                   │ yield status 0  │                  │             │
   │◀── event ─────────│                 │                  │             │
   │                   │ generate_initial_poc               │             │
   │                   │────────────────▶│  _build_prompt   │             │
   │                   │                 │ _call_llm_api_with_settings ──▶│
   │                   │                 │                  │   ◀── JSON ─│
   │                   │                 │ _parse_llm_json_response       │
   │                   │  dict           │                  │             │
   │                   │◀────────────────│                  │             │
   │ yield status 1    │                                                  │
   │◀──────────────────│                                                  │
   │  (若 review=true) │ review_generated_poc                             │
   │                   │────────────────▶│ _build_review_prompt           │
   │                   │                 │ _call_llm_api_with_settings ──▶│
   │                   │                 │  ← JSON ←                      │
   │                   │  dict           │                                │
   │                   │◀────────────────│                                │
   │ yield status 2    │                                                  │
   │◀──────────────────│                                                  │
   │                   │ check_poc_dependencies ──▶ AST/find_spec         │
   │                   │                       ◀── dict                   │
   │                   │ save_poc ────────────────────────▶               │
   │                   │           INSERT poc_records + 写 .py/.json       │
   │                   │◀── poc_id ────────────────                       │
   │ yield result      │                                                  │
   │◀──────────────────│                                                  │
   │ yield [DONE]      │                                                  │
'''.strip('\n'))

    h(doc, '8.2 单条 POC 验证（url_with_params + OOB 不使用）', level=2)
    code(doc, '''
Frontend   FastAPI        PocLibService     DependencyChecker   POC Module
   │ POST /pocs/{id}/execute │                     │                │
   │───────────────────────▶│                     │                │
   │                        │ get_poc_by_id       │                │
   │                        │─── SELECT ─────────────────────▶ SQLite
   │                        │ verification_method=="oob"? 否       │
   │                        │ _execute_python_poc                  │
   │                        │  _normalize_url                      │
   │                        │  check_python_file_dependencies ─────▶│
   │                        │◀────────── ok=true ───────────────────│
   │                        │  动态 importlib + 注入 helper         │
   │                        │  _patch_requests_timeout             │
   │                        │  scan(url, **runtime_params) ──────▶ POC
   │                        │                                  ◀── dict
   │                        │  恢复 requests.Session.request           │
   │                        │  classify_execution_outcome          │
   │ 200 {success, result, classification}                         │
   │◀───────────────────────│                                      │
'''.strip('\n'))

    h(doc, '8.3 批量任务执行', level=2)
    code(doc, '''
Frontend    FastAPI     BatchTaskService       Worker(thread)    PocLibService
   │ POST batch-tasks │                        │                       │
   │────────────────▶│                         │                       │
   │                 │ create_task             │                       │
   │                 │  INSERT batch_tasks + executemany(items)         │
   │                 │ start_task              │                       │
   │                 │  threading.Thread(_run_task) ──▶ 启动           │
   │ 200 {task}      │                         │                       │
   │◀────────────────│                         │                       │
   │                                           │ pending items         │
   │                                           │ ThreadPoolExecutor    │
   │                                           │  submit _execute_task_item
   │                                           │                  ──▶  execute_poc
   │                                           │                  ◀──  outcome
   │                                           │  _store_item_result   │
   │                                           │  _refresh_task_stats  │
   │ GET batch-tasks/{id} (轮询) ──▶            │                       │
   │ ◀── progress ─────                         │                       │
   │                                           │  全部完成 → _finalize │
   │ GET batch-tasks/{id}/export?format=html   │                       │
   │ ◀── 文件 ─────────                         │                       │
'''.strip('\n'))

    h(doc, '8.4 OOB 验证（Interactsh）', level=2)
    code(doc, '''
POC                  oob_service         InteractshClient         Interactsh Server
 │ client=get_oob_client() ──▶ create_client                      │
 │                            ──▶ InteractshClient()              │
 │                                     register (RSA pub)   ────▶│
 │                                          ◀─ aes key + secret ─│
 │ probe=client.build_probe()                                    │
 │    ◀── {url: xxx.interact.sh, flag: xxx}                      │
 │ 发 payload 到目标，包含 probe["url"]                           │
 │ client.verify(flag, protocol="http")                          │
 │    ──▶ poll(secret)                       ────▶               │
 │         ◀── encrypted records ──                              │
 │    RSA decrypt → AES decrypt → 按 flag 过滤                    │
 │    ◀── {hit: bool, records: [...]}                            │
 │ 返回 {"vulnerable": hit, "reason": ..., "details": records}    │
'''.strip('\n'))

    doc.add_page_break()

    # ========== 9 错误处理 ==========
    h(doc, '9 错误处理与日志', level=1)

    h(doc, '9.1 异常分层', level=2)
    table(doc,
        ['层级', '处理策略'],
        [
            ['路由层 (api/routes.py)', '统一 try/except，业务异常 400/404、其它 500；日志记录堆栈'],
            ['服务层 (services/*)', '对外方法重新抛出 Exception 并附加语义信息；对内私有方法保留原始异常'],
            ['POC 执行 (_execute_python_poc)', '捕获所有异常转为 {"success": False, "error": ...} 结构'],
            ['LLM 解析 (_parse_llm_json_response)', '两次解析都失败后才抛异常；失败前保存原始响应至 metadata 目录'],
            ['线程 (_run_task)', 'future 异常统一包装为 {"success": False, "result": {...}}，不让单个子任务错误中断整体调度'],
        ],
        col_widths=[4.5, 11]
    )

    h(doc, '9.2 日志策略', level=2)
    code(doc, '''
logger 配置 (main.py):
  RotatingFileHandler("api_server.log", maxBytes=10MB, backupCount=5)
  StreamHandler (stdout)
  format: %(asctime)s - %(name)s - %(levelname)s - %(message)s

HTTP 中间件:
  每个请求记录:
    "收到请求: {method} {url}"
    "客户端: {client_host}"
    "响应状态: {code} | 耗时: {t:.3f}秒"

关键业务日志:
  LLM 调用前后：api base / model / temperature / prompt 长度
  批量任务：创建/启动/子任务执行/完成/取消
  POC 生成结果保存：poc_id 与简要说明
  OOB 配置变更、空间测绘调用
'''.strip('\n'))

    h(doc, '9.3 敏感信息脱敏规则', level=2)
    bullet(doc, 'API Key：保留前 7 位 + 尾 4 位，中间以 "..." 替代；长度 ≤ 10 时仅保留前 3 位 + "***"')
    bullet(doc, 'OOB Token / 空间测绘 Token：保留前 3 位 + 尾 3 位，中间 "..."')
    bullet(doc, '配置文件直接以明文保存在本地 JSON，部署者需要保证文件权限；前端 /api/config/* 响应始终是脱敏版本')

    h(doc, '9.4 静态资源缓存控制', level=2)
    para(doc, '开发期为减少"改完代码忘了刷新"的心智负担，主页面与 Nuclei 页面响应增加：Cache-Control: no-store, no-cache, must-revalidate, max-age=0；Pragma: no-cache；Expires: 0。生产部署可按需放开。')

    doc.add_page_break()

    # ========== 10 性能与并发 ==========
    h(doc, '10 性能与并发设计', level=1)

    h(doc, '10.1 并发模型', level=2)
    bullet(doc, 'HTTP 层：FastAPI + Uvicorn 单进程；默认使用 asyncio 事件循环，LLM 调用通过 AsyncOpenAI 非阻塞')
    bullet(doc, '批量任务：每个任务一个后台 Thread（守护线程）+ ThreadPoolExecutor(max_workers=concurrency) 负责子任务；concurrency ∈ [1, 5]')
    bullet(doc, 'Nuclei 流式扫描：asyncio.subprocess，避免阻塞主事件循环')
    bullet(doc, 'POC 动态执行：同步阻塞调用，但通过 requests 超时补丁限制单次请求 ≤ 8s，避免拖垮线程')

    h(doc, '10.2 性能敏感点与优化手段', level=2)
    table(doc,
        ['敏感点', '现象', '当前策略', '后续优化建议'],
        [
            ['Nuclei 模板加载', '模板数 10k+ 时扫描目录 1~2s', '_parse_template_fast + 内存缓存 + cache/clear 接口', '按需懒加载文件夹子项；文件 mtime 做缓存失效'],
            ['LLM 调用慢', '首次冷启动 + 长 Prompt 可能 30s+', '异步 SDK + SSE 过程推送，让用户感知进度', '可增加流式 token 输出（目前整体 JSON 返回）'],
            ['批量任务大规模扫描', '笛卡尔积可达 2000 子任务', 'MAX_* 限制 + 并发上限 5', '迁移到 Redis + Celery / ARQ 异步队列'],
            ['SQLite 写锁', '线程池高并发 UPDATE 可能短暂阻塞', 'contextmanager 快速释放连接', '启用 WAL 模式；或按任务分文件库'],
            ['POC 动态导入', '每次生成新模块名，sys.modules 暂存', 'finally 清理；helper 模块同样每次注入 / 清理', '若 POC 文件名已存在可考虑 importlib.reload 缓存'],
        ],
        col_widths=[3, 3.5, 4.5, 4.5]
    )

    h(doc, '10.3 资源限制', level=2)
    code(doc, '''
BatchTaskService:
    MAX_URLS             = 200
    MAX_POCS             = 50
    MAX_TASK_ITEMS       = 2000
    DEFAULT_CONCURRENCY  = 3
    MAX_CONCURRENCY      = 5

HTTPRuntimeClient:
    timeout              = 6  (秒)

_patch_requests_timeout:
    default              = 8  (秒)

OOBClient:
    poll_interval        = 1.0 (用户可配)
    max_polls            = 3   (用户可配)

Nuclei scan:
    single / multiple    = 60~120 秒
    folder               = 300 秒
'''.strip('\n'))

    doc.add_page_break()

    # ========== 11 测试设计 ==========
    h(doc, '11 测试设计', level=1)

    h(doc, '11.1 单元测试要点', level=2)
    table(doc,
        ['模块', '测试用例', '断言重点'],
        [
            ['_parse_llm_json_response', '合法 JSON / 带 ```json 包裹 / 含未转义换行 / 完全损坏', '前三项应成功；第四项抛异常并生成 last_llm_response.json'],
            ['_normalize_poc_code_string', '含字面量 \\n 的单行代码 / 正常多行代码 / 空字符串', '仅对第一种转为真实换行，其余保持不变'],
            ['classify_execution_outcome', '成功 / vuln=false / 网络超时 / 缺依赖 / nuclei timeout / oob 未配置', '返回正确的 category/code/stage/retryable'],
            ['check_python_code_dependencies', '纯标准库 / requests / 不存在的模块 / 相对导入 / 语法错误', 'ok 与 missing 列表符合预期；语法错误返回 parse_error'],
            ['BatchTaskService.create_task', '空列表 / 超过 MAX_* / 合法任务', '前两种 ValueError；合法任务产出正确 total_items'],
            ['_run_task 取消', '启动后立即 cancel_task', '所有未开始子任务落为 cancelled；任务 status=cancelled'],
            ['OOB 解析', '模拟 interactsh 密文 / ceye JSON 响应', 'verify 返回 hit/records 与模拟一致'],
        ],
        col_widths=[4, 5.5, 6]
    )

    h(doc, '11.2 接口集成测试', level=2)
    bullet(doc, 'POST /api/generate-poc：mock llm_service.generate_initial_poc，验证 SSE 事件顺序与 result 结构')
    bullet(doc, 'GET /api/pocs/search：准备测试数据后按 vuln_type、keyword、verifiable、limit/offset 组合验证')
    bullet(doc, 'POST /api/batch-tasks：分别提交 1×1、N×M、超限三种样例，检查数据库 items 数量与 400 错误')
    bullet(doc, 'GET /api/batch-tasks/{id}/export：html/json/txt 三种格式都能正确响应 Content-Disposition')
    bullet(doc, 'POST /api/pocs/{id}/execute：覆盖 url_only、url_with_params、verification_method=oob 未配置三种路径')
    bullet(doc, 'GET/POST /api/config/*：更新后 GET 返回脱敏值，重启服务后配置仍存在')

    h(doc, '11.3 端到端手工验收', level=2)
    bullet(doc, '从生成页输入一段 SQLi 描述 → 开启二次审核 → 观察前端进度 → POC 保存至库')
    bullet(doc, '在 POC 库选中一条 url_with_params 类型 → 检查参数表单渲染 → 执行 → 验证结果')
    bullet(doc, '批量任务：2 URL × 3 POC，观察进度 → 中途取消 → 剩余 items 标记 cancelled')
    bullet(doc, 'Nuclei 页面：选择 http/cves 目录 → 勾选 3 个模板 → 对 http://example.com 扫描')
    bullet(doc, '系统配置：依次修改 LLM、审核模型、OOB、FOFA → 重启后检查配置仍生效')

    doc.add_page_break()

    # ========== 12 附录 ==========
    h(doc, '12 附录', level=1)

    h(doc, '12.1 execution_mode × verification_method 组合矩阵', level=2)
    table(doc,
        ['execution_mode', 'verification_method', 'input_schema', 'manual_steps', '运行时行为'],
        [
            ['url_only', 'direct', '不使用', '不使用', 'scan(url)；POC 内部自行判断'],
            ['url_only', 'oob', '不使用', '不使用', 'OOB 预检 → scan(url) 内部使用 helper verify'],
            ['url_with_params', 'direct', '必填', '不使用', '前端按 schema 弹表单 → scan(url, **params)'],
            ['url_with_params', 'oob', '必填', '不使用', '同上 + OOB 预检'],
            ['manual_guide', 'manual', '不使用', '必填', '前端展示步骤；不执行 scan'],
        ],
        col_widths=[3, 3, 2.5, 2.5, 4.5]
    )

    h(doc, '12.2 SSE 事件结构汇总', level=2)
    code(doc, '''
POC 生成 (POST /api/generate-poc):
  data: {"type":"status", "step": 0|1|2, "message": str, "status": "active|completed"}
  data: {"type":"result", "data": PocResponse}
  data: {"type":"error",  "data": PocResponse(success=false)}
  data: [DONE]

Nuclei 流式 (POST /api/nuclei/scan/stream):
  data: {"type":"log",     "message": str}
  data: {"type":"finding", "data": NucleiFinding}
  data: {"type":"done",    "summary": {total_findings, vulnerable}}
  data: {"type":"error",   "message": str}
  data: [DONE]
'''.strip('\n'))

    h(doc, '12.3 依赖清单说明', level=2)
    table(doc,
        ['依赖', '版本约束', '用途'],
        [
            ['fastapi', '>=0.110.0', 'Web 框架 + OpenAPI 文档'],
            ['uvicorn[standard]', '>=0.27.0', 'ASGI 服务器'],
            ['pydantic', '>=2.6.0', '请求/响应模型'],
            ['openai', '>=1.0.0', '兼容 OpenAI 协议的 LLM 客户端'],
            ['PyYAML', '>=6.0.1', 'Prompt 模板 + Nuclei 模板解析'],
            ['dnspython', '>=2.4.0', 'OOB DNS 查询兜底'],
            ['requests', '>=2.31.0', 'POC 内 HTTP 调用'],
            ['pycryptodome', '>=3.20.0', 'Interactsh AES 解密'],
            ['itsdangerous', '>=2.2.0', '敏感配置签名'],
            ['Flask', '>=3.0.0', '兼容少量旧脚本'],
        ],
        col_widths=[3.5, 2.5, 9.5]
    )

    h(doc, '12.4 可扩展点', level=2)
    bullet(doc, 'LLM Provider：通过 update_config 更换 base_url 即可切换；若需要原生 API（非 OpenAI 兼容），在 _call_llm_api_with_settings 中分派')
    bullet(doc, 'OOB Provider：向 OOBService.provider_factories 注册新 Provider 即可；BaseOOBClient 接口不变')
    bullet(doc, '空间测绘 Provider：在 asset_source_service 内新增查询函数并加入分派表')
    bullet(doc, '批量任务 Engine：目前 engine_type = "poc" 或 "nuclei"；新增 Engine 时只需扩展 _execute_task_item 分派与 batch_task_items.engine_type 取值')
    bullet(doc, '结果分类：failure_classifier 为关键字匹配，可接入更细粒度的规则或基于 AI 的分类')

    h(doc, '12.5 文档变更记录', level=2)
    table(doc,
        ['版本', '日期', '变更说明', '作者'],
        [
            ['V1.0', '2026-04-18', '初版：基于当前代码库完成全量详细设计', '项目组'],
        ],
        col_widths=[2, 3, 9, 2]
    )

    # 保存
    doc.save(str(OUTPUT_PATH))
    print(f'文档已生成: {OUTPUT_PATH}')
    print(f'文件大小: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB')


if __name__ == '__main__':
    build()
