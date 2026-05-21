"""
AI-POC 概要设计文档生成脚本
使用 python-docx 生成标准格式的 Word 文档
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = Path(__file__).parent / "AI-POC概要设计文档.docx"


# ================= 样式辅助 =================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), 'single')
            tag.set(qn('w:sz'), '4')
            tag.set(qn('w:color'), '808080')
            tcBorders.append(tag)


def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, name='黑体', size=sizes.get(level, 12), bold=True, color=(0x1F, 0x3A, 0x5F))
    return p


def add_para(doc, text, indent_first=True, bold=False, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_font(run, name='宋体', size=size, bold=bold)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    run = p.runs[0] if p.runs else p.add_run()
    # 重置内容
    for r in p.runs:
        r.text = ''
    run = p.add_run(text)
    set_font(run, name='宋体', size=11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, name='Consolas', size=10, color=(0x2F, 0x4F, 0x4F))
    # 背景灰
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F5F5F5')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, name='黑体', size=11, bold=True, color=(0xFF, 0xFF, 0xFF))
        # 表头背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F3A5F')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        set_cell_border(hdr_cells[i], top=True, bottom=True, left=True, right=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run(str(val))
            set_font(run, name='宋体', size=10)
            set_cell_border(row_cells[i], top=True, bottom=True, left=True, right=True)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ================= 构建文档 =================

def build():
    doc = Document()

    # 基础样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr() if hasattr(style.element, 'get_or_add_rPr') else None

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # ---------- 封面 ----------
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI-POC 漏洞验证平台')
    set_font(run, name='黑体', size=32, bold=True, color=(0x1F, 0x3A, 0x5F))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('概要设计文档')
    set_font(run, name='黑体', size=26, bold=True, color=(0x1F, 0x3A, 0x5F))

    for _ in range(8):
        doc.add_paragraph()

    meta_rows = [
        ('文档名称', 'AI-POC 漏洞验证平台概要设计文档'),
        ('文档版本', 'V1.0'),
        ('编写日期', '2026-04-18'),
        ('文档状态', '初稿'),
        ('适用范围', '软件工程课程设计 / 项目答辩'),
    ]
    t = doc.add_table(rows=len(meta_rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, (k, v) in enumerate(meta_rows):
        row = t.rows[i]
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
        for j, text in enumerate((k, v)):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, name='宋体', size=12, bold=(j == 0))
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    # ---------- 目录占位 ----------
    add_heading(doc, '目  录', level=1)
    toc_items = [
        '1 引言', '2 总体设计', '3 系统架构设计', '4 功能模块设计',
        '5 数据库设计', '6 接口设计', '7 关键流程设计',
        '8 安全设计', '9 运行部署设计', '10 技术难点与解决方案', '11 附录',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.6
        run = p.add_run(item)
        set_font(run, name='宋体', size=12)
    add_para(doc, '注：打开文档后可按 Ctrl+A 全选 → 按 F9 刷新自动生成的编号或页码；本目录页供阅读导航使用。', indent_first=False, size=10)

    doc.add_page_break()

    # ========== 1 引言 ==========
    add_heading(doc, '1 引言', level=1)

    add_heading(doc, '1.1 编写目的', level=2)
    add_para(doc, '本文档是 AI-POC 漏洞验证平台的概要设计文档，主要用于在需求分析阶段成果的基础上，明确系统的总体架构、模块划分、数据结构、接口设计与关键业务流程，作为后续详细设计、编码实现、测试验证的依据。本文档的读者对象包括项目指导教师、课程答辩评审老师、项目开发成员以及后续可能接手维护的同学。')

    add_heading(doc, '1.2 项目背景', level=2)
    add_para(doc, '当前 Web 漏洞验证工作中，POC（Proof of Concept）编写、调试、管理和复用往往依赖人工完成，存在生成效率低、验证条件不统一、历史结果难追踪等问题。随着大语言模型能力的提升，利用 LLM 辅助生成验证型 POC 已具备可行性。')
    add_para(doc, 'AI-POC 项目面向这一场景，尝试构建一个以 AI 生成 POC 为核心、兼容多种验证方式的漏洞验证平台，实现从漏洞信息输入、POC 自动生成、验证执行到检测记录沉淀的完整流程，同时兼容 Nuclei 等现有扫描器模板执行能力。')

    add_heading(doc, '1.3 术语定义', level=2)
    add_table(doc,
        ['术语/缩写', '说明'],
        [
            ['POC', 'Proof of Concept，漏洞概念验证代码或步骤，用于证明漏洞存在'],
            ['LLM', 'Large Language Model，大语言模型'],
            ['OOB', 'Out-of-Band，带外通信，用于验证无回显类漏洞（如盲注、SSRF）'],
            ['Nuclei', '开源的基于 YAML 模板的漏洞扫描器'],
            ['Interactsh', '开源的 OOB 监听服务，可自建或使用公开服务'],
            ['CEye', '常用的 DNS/HTTP 外带检测服务'],
            ['SSE', 'Server-Sent Events，服务端向浏览器单向流式推送事件'],
            ['POC 库', '本系统用于分类、存储、检索、执行 POC 的数据与文件集合'],
            ['空间测绘平台', '如 FOFA、Hunter、Quake 等支持批量查询互联网资产的平台'],
        ],
        col_widths=[3.5, 12]
    )

    add_heading(doc, '1.4 参考资料', level=2)
    add_bullet(doc, '《AI-POC 项目开题文档》（本项目需求分析部分）')
    add_bullet(doc, 'FastAPI 官方文档（https://fastapi.tiangolo.com/）')
    add_bullet(doc, 'OpenAI Python SDK 文档')
    add_bullet(doc, 'Nuclei 官方文档（https://docs.projectdiscovery.io/tools/nuclei）')
    add_bullet(doc, 'Interactsh 项目说明文档')
    add_bullet(doc, 'SQLite 官方文档（https://www.sqlite.org/docs.html）')

    doc.add_page_break()

    # ========== 2 总体设计 ==========
    add_heading(doc, '2 总体设计', level=1)

    add_heading(doc, '2.1 设计目标', level=2)
    add_para(doc, '系统总体设计目标可以归纳为四个方面：一是降低 POC 编写成本，借助大模型对漏洞信息的理解能力自动生成验证代码或人工指南；二是统一 POC 管理入口，将 AI 生成结果、Nuclei 模板以及人工维护的 POC 纳入同一套管理体系；三是提供多种执行形态，覆盖可直接验证、带参数验证、外带验证和人工指南等场景；四是支持任务化检测与结果沉淀，便于课程演示和后续分析复用。')

    add_heading(doc, '2.2 设计原则', level=2)
    add_bullet(doc, '模块化：接口层、业务层、存储层、前端层分离，各业务服务单一职责')
    add_bullet(doc, '可扩展：LLM、OOB、空间测绘平台均通过 Provider 抽象，支持多供应商切换')
    add_bullet(doc, '稳定优先：大模型输出结果经过依赖预检、二次审核、失败分类等防线保护')
    add_bullet(doc, '轻量化：使用 SQLite + 文件目录混合存储，无需额外部署数据库中间件')
    add_bullet(doc, '可演示：支持本地便携部署，提供直观的前端操作界面')
    add_bullet(doc, '安全合规：所有验证能力默认提示仅限授权环境使用，敏感配置加密保存')

    add_heading(doc, '2.3 系统功能总览', level=2)
    add_para(doc, 'AI-POC 平台从功能角度可以划分为六大子系统：POC 智能生成、POC 库管理、单条验证执行、批量检测任务、Nuclei 兼容扫描、系统配置管理。六大子系统围绕"POC 的全生命周期"展开，形成从生成到归档的闭环。')

    add_table(doc,
        ['子系统', '核心功能', '典型使用场景'],
        [
            ['POC 智能生成', '基于漏洞描述/CVE/HTTP 报文调用 LLM 生成 POC 初稿，可启用二次审核', '新漏洞快速产出可执行 POC'],
            ['POC 库管理', 'POC 分类展示、搜索过滤、查看代码、删除、统计', '历史 POC 的复用与管理'],
            ['单条验证执行', '对指定 POC 输入目标 URL 和参数后执行，返回结构化结果', '单点验证与调试'],
            ['批量检测任务', '对 URL 集合 × POC 集合发起异步任务，支持取消和导出', '成片资产快速过检'],
            ['Nuclei 兼容扫描', '封装 Nuclei CLI，支持模板浏览、即时扫描和任务化扫描', '借助社区模板提高覆盖面'],
            ['系统配置管理', '维护 LLM、二次审核模型、OOB、空间测绘平台等配置', '部署与日常运维'],
        ],
        col_widths=[3.5, 7.5, 5]
    )

    add_heading(doc, '2.4 运行环境', level=2)
    add_table(doc,
        ['项目', '配置要求', '说明'],
        [
            ['操作系统', 'Windows 10/11，兼容 Linux/macOS', '当前主要针对 Windows 本地环境'],
            ['Python 版本', '3.11 及以上', '依赖 FastAPI、Pydantic v2'],
            ['Web 框架', 'FastAPI + Uvicorn', '提供 REST API 与 SSE 流式接口'],
            ['数据库', 'SQLite 3（文件型）', 'POC 元数据、批量任务元数据'],
            ['前端技术', '原生 HTML + CSS + JavaScript', '无构建工具，可直接打开'],
            ['Nuclei', '官方二进制可执行文件', '模板扫描能力（可选）'],
            ['LLM 服务', '支持 OpenAI 兼容协议', '可对接 OpenAI、DeepSeek、Qwen、本地模型等'],
            ['浏览器', 'Chrome / Edge 近 2 年版本', '前端依赖 ES6+ 特性'],
        ],
        col_widths=[3, 5, 7.5]
    )

    doc.add_page_break()

    # ========== 3 系统架构设计 ==========
    add_heading(doc, '3 系统架构设计', level=1)

    add_heading(doc, '3.1 总体架构', level=2)
    add_para(doc, '系统采用"前端 + API + 服务层 + 存储层 + 外部能力层"的五层架构。前端通过浏览器加载的静态页面与用户交互，API 层基于 FastAPI 定义统一入口，业务服务层将功能拆分为多个独立 Service，存储层由 SQLite 数据库与 POC 文件目录共同组成，外部能力层负责对接 LLM、OOB、Nuclei 与空间测绘平台。')

    add_code(doc, '''
┌──────────────────────────────────────────────────────────────┐
│                       前端静态页面层                          │
│  index.html (生成/POC库/检测记录)   nuclei.html (模板扫描)    │
│  app.js / library.js / nuclei.js      *.css                   │
└───────────────┬──────────────────────────────────────────────┘
                │  HTTP / SSE
┌───────────────▼──────────────────────────────────────────────┐
│                    FastAPI 接口路由层 (api/routes.py)         │
│  POC 生成 | POC 库 | 批量任务 | Nuclei | LLM 配置 | OOB | 资产 │
└───────────────┬──────────────────────────────────────────────┘
                │
┌───────────────▼──────────────────────────────────────────────┐
│                      业务服务层 (services/)                   │
│ llm_service / poc_library_service / batch_task_service /      │
│ nuclei_service / oob_service / http_runtime /                 │
│ asset_source_service / dependency_checker / failure_classifier│
└─────┬───────────────┬───────────────┬──────────────┬─────────┘
      │               │               │              │
┌─────▼──────┐ ┌──────▼─────┐ ┌───────▼─────┐ ┌──────▼───────┐
│ SQLite     │ │ POC 文件    │ │ Nuclei CLI  │ │ LLM / OOB /  │
│ 数据库     │ │ 目录 pocs/  │ │ 与模板库    │ │ 测绘平台     │
└────────────┘ └────────────┘ └─────────────┘ └──────────────┘
'''.strip('\n'))

    add_heading(doc, '3.2 技术栈选型', level=2)
    add_table(doc,
        ['分层', '技术选型', '选型依据'],
        [
            ['前端', '原生 HTML + CSS + 原生 JS（含 Fetch + EventSource）', '零构建，部署简单，便于课程演示；满足界面需求即可，不需要前端工程化'],
            ['后端框架', 'FastAPI + Uvicorn', '异步支持良好，自动生成 OpenAPI 文档，Pydantic 校验请求体'],
            ['数据校验', 'Pydantic v2', '与 FastAPI 深度集成，模型即文档'],
            ['LLM 客户端', 'openai Python SDK（兼容 OpenAI 协议）', '通过 base_url 切换多家供应商，复用同一套接口'],
            ['OOB 对接', 'Interactsh 自建协议 + CEye HTTP API', '兼顾开源自建和云端服务，覆盖主流外带验证场景'],
            ['数据库', 'SQLite 3', '文件型零部署，适合本地演示；结构稳定便于后期迁移'],
            ['加密', 'pycryptodome（AES）+ itsdangerous', 'Interactsh 通信解密与配置 Token 加密'],
            ['DNS 查询', 'dnspython', 'OOB DNS 记录查询兜底'],
            ['扫描兼容', 'Nuclei 官方二进制', '复用成熟社区模板'],
            ['日志', 'logging + RotatingFileHandler', '轻量，10MB 轮转，5 份备份'],
        ],
        col_widths=[2.5, 5.5, 7]
    )

    add_heading(doc, '3.3 部署架构', level=2)
    add_para(doc, '系统以单节点本地部署为主。Uvicorn 在 127.0.0.1:8000 启动 FastAPI 应用，浏览器通过 HTTP 访问前端页面，并通过同域的 REST 接口与后台交互。系统对外依赖的服务按需调用：LLM 调用通过 HTTPS 访问模型供应商；OOB 调用根据 provider 类型访问 Interactsh 服务或 CEye API；Nuclei 作为本地子进程被 nuclei_service 通过命令行调起；空间测绘平台则在资产导入时按需查询。')

    add_code(doc, '''
[浏览器]  ──HTTP──▶  [FastAPI @ 127.0.0.1:8000]
                               │
                               ├──▶ SQLite (pocs/poc_library.db)
                               ├──▶ POC 文件 (pocs/python/*.py)
                               ├──▶ Nuclei.exe + 模板 (pocs/nuclei/)
                               ├──▶ LLM API (HTTPS)
                               ├──▶ Interactsh / CEye (HTTPS)
                               └──▶ FOFA / Hunter / Quake (HTTPS)
'''.strip('\n'))

    add_heading(doc, '3.4 目录结构', level=2)
    add_table(doc,
        ['顶层目录 / 文件', '作用说明'],
        [
            ['main.py', 'FastAPI 应用入口，负责装配路由、日志、静态资源挂载'],
            ['config.py', '全局配置（API 主机端口、安全警告语）'],
            ['requirements.txt', 'Python 依赖清单'],
            ['nuclei.exe', 'Windows 下 Nuclei 可执行文件'],
            ['api/', '路由层：routes.py 集中声明所有 REST 端点'],
            ['services/', '业务服务层：9 个独立 Service 模块'],
            ['models/', '数据模型：schemas.py 声明所有 Pydantic 请求/响应模型'],
            ['frontend/', '前端静态资源：index.html / nuclei.html / css/ / js/'],
            ['prompts/', 'LLM 提示词模板：poc_generation.yaml'],
            ['pocs/', 'POC 数据根目录：python/（代码）、metadata/（元数据 JSON）、nuclei/（模板）、poc_library.db（数据库）'],
            ['tests/', '自动化测试代码'],
            ['docs/', '项目文档：analysis/（分析与设计）'],
            ['release/', '便携版打包产物'],
            ['api_server.log', '应用运行日志（自动轮转）'],
        ],
        col_widths=[4.5, 11]
    )

    doc.add_page_break()

    # ========== 4 功能模块设计 ==========
    add_heading(doc, '4 功能模块设计', level=1)

    add_heading(doc, '4.1 模块划分总览', level=2)
    add_para(doc, '按照前后端分层结构，系统拆分为前端 UI 模块、API 路由模块、业务服务模块、数据模型模块、外部能力接入模块共五大类。下文按模块给出职责、核心类与关键方法说明，以便详细设计阶段直接对照实现。')

    add_heading(doc, '4.2 API 路由模块（api/routes.py）', level=2)
    add_para(doc, 'routes.py 是系统唯一的 REST 入口，所有业务请求均经此转发至对应 Service。路由按功能分组，使用 FastAPI 的装饰器声明 HTTP 方法、路径、摘要与响应模型，并通过依赖注入自动完成请求体校验。下表给出主要端点分组：')
    add_table(doc,
        ['功能分组', '端点示例', '对应服务'],
        [
            ['POC 生成', 'POST /api/generate-poc （SSE 流式）', 'llm_service + poc_library_service'],
            ['POC 库', 'GET /api/pocs/search、GET /api/pocs/{id}、POST /api/pocs/{id}/execute、DELETE /api/pocs/{id}、GET /api/pocs/statistics、GET /api/pocs/vuln-types、GET /api/pocs/{id}/code', 'poc_library_service'],
            ['批量任务', 'POST /api/batch-tasks、GET /api/batch-tasks、GET /api/batch-tasks/{id}、GET /api/batch-tasks/{id}/items、POST /api/batch-tasks/{id}/cancel、GET /api/batch-tasks/{id}/export', 'batch_task_service'],
            ['Nuclei', 'GET /api/nuclei/status、GET /api/nuclei/folders、GET /api/nuclei/templates、GET /api/nuclei/template/content、POST /api/nuclei/scan、POST /api/nuclei/scan/stream、POST /api/nuclei/tasks、POST /api/nuclei/cache/clear', 'nuclei_service + batch_task_service'],
            ['LLM 配置', 'GET/POST /api/config/llm、GET/POST /api/config/review-llm', 'llm_service'],
            ['OOB 配置', 'GET/POST /api/config/oob', 'oob_service'],
            ['空间测绘', 'GET/POST /api/config/asset-sources、POST /api/asset-sources/import', 'asset_source_service'],
        ],
        col_widths=[2.5, 9, 4]
    )

    add_heading(doc, '4.3 业务服务模块（services/）', level=2)

    add_heading(doc, '4.3.1 llm_service.py — 大模型生成服务', level=3)
    add_para(doc, '该服务封装对大语言模型的访问，核心类 LLMService 负责配置管理、Prompt 组装、API 调用、响应解析与二次审核。关键方法包括：generate_initial_poc(vulnerability_info, target_info) 生成初稿；review_generated_poc(...) 基于审核模型对初稿进行修正；_build_prompt() 和 _build_review_prompt() 分别组装两阶段 Prompt；_call_llm_api_with_settings() 使用 openai SDK 发起 HTTP 请求；_parse_llm_json_response() 处理大模型返回中常见的控制字符、转义错误等鲁棒性问题；update_config()/update_review_config() 将用户修改的配置同步到 llm_config.json 持久化。服务通过 prompts/poc_generation.yaml 加载系统提示词模板，便于脱离代码调优 Prompt。')

    add_heading(doc, '4.3.2 poc_library_service.py — POC 库管理服务', level=3)
    add_para(doc, '核心类 PocLibraryService 同时管理数据库元数据与磁盘 POC 文件。init_database() 在启动时自动建表并做字段兼容迁移；save_poc() 将 LLM 生成结果写入 poc_records 表并保存脚本文件到 pocs/python/ 目录；get_poc_by_id()/search_pocs() 提供基础查询；execute_poc() 是统一执行入口，根据 poc_type 与 execution_mode 分发给 _execute_python_poc() 或 Nuclei 执行器，并通过 _invoke_scan() 动态加载并调用 POC 文件中的 scan(url, **kwargs) 函数；check_poc_dependencies() 调用 dependency_checker 做静态依赖预检；delete_poc() 负责元数据与文件的一并删除。')

    add_heading(doc, '4.3.3 batch_task_service.py — 批量任务执行服务', level=3)
    add_para(doc, 'BatchTaskService 负责批量检测任务的全生命周期。create_task() 对 URL 与 POC 做去重、限流（MAX_URLS 与 MAX_POCS）与合法性校验，写入 batch_tasks 主表并展开笛卡尔积写入 batch_task_items 明细；start_task() 以线程池方式并发执行任务，通过 cancel_event 支持用户主动取消；_execute_task_item() 为 POC 任务，调用 poc_library_service.execute_poc()；_execute_nuclei_task_item() 为 Nuclei 任务，调用 nuclei_service.scan_single()；_refresh_task_stats() 实时汇总任务总量、成功数、失败数与漏洞命中数；export_task_report() 支持 HTML / JSON / TXT 三种格式报告导出；另外还封装了子任务结果摘要的落盘与回填（_write_detail_file、_backfill_batch_task_item_summaries）。')

    add_heading(doc, '4.3.4 nuclei_service.py — Nuclei 扫描封装', level=3)
    add_para(doc, 'NucleiService 将 Nuclei 命令行封装为可被 API 调用的服务。check_nuclei_available() 执行 nuclei -version 验证可用性；get_folder_structure() 扫描 pocs/nuclei/templates 目录并统计每个子目录的模板数量；get_templates_by_folder() 对模板做分页与关键字过滤，内部使用 _parse_template_fast() 快速解析 YAML 头部字段（id、name、author、severity、tags 等）；scan_single()/scan_multiple()/scan_folder() 提供三种粒度的同步扫描；scan_stream_async() 基于 subprocess 异步管道以 SSE 形式返回实时发现；同时带有模板解析缓存以降低大量 YAML 扫描的 I/O 压力。')

    add_heading(doc, '4.3.5 oob_service.py — OOB 外带验证服务', level=3)
    add_para(doc, 'OOBService 统一管理外带验证能力，通过 Provider 抽象支持 Interactsh 自建协议和 CEye 公共 API。BaseOOBClient 定义 build_probe(protocol, length, value) 与 verify(flag) 两个核心接口；InteractshClient 使用 RSA 注册 + AES 解密机制获取命中记录；CEyeClient 直接通过 HTTP API 查询 DNS/HTTP 日志。POC 内部可以通过调用 oob_service.build_probe() 获取一次性回连地址，并在执行结束后通过 verify() 判断目标是否发起了外带请求，从而支撑 SSRF、盲注、命令执行等无回显类漏洞的验证。')

    add_heading(doc, '4.3.6 http_runtime.py — HTTP 执行运行时', level=3)
    add_para(doc, 'http_runtime 提供统一的 HTTP 请求封装与超时补丁。_patch_requests_timeout() 为 requests 库的会话方法注入默认超时，避免 POC 因目标无响应而长时间阻塞整个批量任务。模块同时提供 URL 归一化、重定向控制等工具函数，保证不同 POC 在同一运行时下具有一致的网络行为。')

    add_heading(doc, '4.3.7 asset_source_service.py — 空间测绘平台对接', level=3)
    add_para(doc, 'AssetSourceService 抽象 FOFA / Hunter / Quake 三类平台，统一了 provider 配置、查询语法与分页参数。update_provider_config() 持久化平台凭证；import_targets() 根据 provider 选择对应 SDK 或 HTTP API 查询目标，随后将结果标准化为 URL 列表返回给前端，直接用于批量任务的目标输入。')

    add_heading(doc, '4.3.8 dependency_checker.py / failure_classifier.py — 质量保障', level=3)
    add_para(doc, 'dependency_checker 对 LLM 生成的 POC 代码做静态 AST 分析，提取 import 语句并与白名单对比，提前发现"目标环境缺少某库"导致运行时报错的问题。failure_classifier 则对执行失败结果进行归因，分为网络类（连接超时、DNS 失败）、环境类（依赖缺失、Python 语法错误）、代码类（POC 内部异常）、OOB 类（未命中回连）等，便于用户在检测记录页快速定位问题来源，而不是仅看到一句"失败"。')

    add_heading(doc, '4.4 数据模型模块（models/schemas.py）', level=2)
    add_para(doc, '所有请求与响应均使用 Pydantic BaseModel 定义，避免裸 dict 流转。主要模型分组如下：')
    add_table(doc,
        ['模型类别', '代表类', '关键字段'],
        [
            ['POC 生成', 'VulnerabilityRequest / PocResponse', 'vulnerability_info、enable_second_review、execution_mode、verification_method、input_schema、manual_steps、dependency_check'],
            ['POC 执行', 'ScanRequest', 'target_url、runtime_params'],
            ['LLM 配置', 'LLMConfigRequest / LLMConfigResponse', 'api_key、model_id、base_url、temperature、max_tokens'],
            ['OOB 配置', 'OOBConfigRequest / OOBConfigResponse', 'enabled、provider、interactsh_server、interactsh_token、ceye_token、poll_interval、max_polls'],
            ['空间测绘', 'AssetSourceConfigRequest / AssetSourceImportRequest', 'provider、email、token、query、pages'],
            ['Nuclei', 'NucleiScanRequest / NucleiTaskCreateRequest / NucleiTemplate / NucleiFinding / NucleiScanResponse / NucleiStatusResponse', 'target_url、template_paths、folder、timeout、severity、findings'],
            ['批量任务', 'BatchTaskCreateRequest / BatchTaskActionResponse', 'target_urls、poc_ids、concurrency、task'],
        ],
        col_widths=[2.5, 5, 8]
    )

    add_heading(doc, '4.5 前端模块（frontend/）', level=2)
    add_para(doc, '前端为静态页面，无任何构建工具依赖。主要由两个 HTML 入口、三个 JS 模块和三个样式文件组成：')
    add_table(doc,
        ['文件', '职责'],
        [
            ['index.html', '系统主页面，承载 POC 生成、POC 库、检测记录、系统配置四大面板的 Tab 切换'],
            ['nuclei.html', 'Nuclei 扫描器独立页面，承载模板浏览、单目标扫描、批量任务创建'],
            ['js/app.js', '生成页逻辑：表单提交、EventSource 接入流式进度、结果渲染、配置弹窗'],
            ['js/library.js', 'POC 库逻辑：分类筛选、分页搜索、代码查看、单次执行、批量任务创建、检测记录轮询与报告导出'],
            ['js/nuclei.js', 'Nuclei 页面逻辑：文件夹树、模板列表分页、流式扫描、任务提交'],
            ['css/style.css', '主页面整体样式：布局、Tab、表单、按钮、通用组件'],
            ['css/new-library.css', 'POC 库及检测记录的专用样式'],
            ['css/nuclei.css', 'Nuclei 页面的专用样式'],
        ],
        col_widths=[3.5, 12]
    )

    doc.add_page_break()

    # ========== 5 数据库设计 ==========
    add_heading(doc, '5 数据库设计', level=1)

    add_heading(doc, '5.1 存储方案', level=2)
    add_para(doc, '系统采用"SQLite 元数据库 + 文件系统代码库"的混合存储方案。结构化信息（POC 元数据、批量任务、子任务执行结果）保存在 SQLite 中，便于检索与统计；而 POC 代码、元数据 JSON、Nuclei 模板、批量任务详情 JSON 等大文本与二进制内容则以文件形式落盘，避免数据库体积失控，同时便于用户直接查看、修改、备份。')

    add_heading(doc, '5.2 数据库表结构', level=2)

    add_heading(doc, '5.2.1 poc_records — POC 元数据表', level=3)
    add_para(doc, '记录每个 POC 的基本信息、执行模式、验证方式、文件路径与统计信息。')
    add_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK AUTOINCREMENT', '主键'],
            ['vuln_type', 'TEXT', 'NOT NULL', '漏洞类型（sqli / xss / rce 等）'],
            ['vuln_name', 'TEXT', 'NOT NULL', '漏洞名称，含产品或 CVE 编号'],
            ['vuln_description', 'TEXT', '', '用户输入的原始漏洞信息'],
            ['poc_type', 'TEXT', "DEFAULT 'python'", 'python / nuclei / manual'],
            ['poc_file_path', 'TEXT', 'NOT NULL', 'POC 脚本或模板文件绝对路径'],
            ['create_time', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
            ['last_used', 'TIMESTAMP', '', '最近一次执行时间'],
            ['tags', 'TEXT', '', 'JSON 数组，标签'],
            ['metadata', 'TEXT', '', 'JSON，扩展元数据（如目标信息）'],
            ['verifiable', 'BOOLEAN', 'DEFAULT 1', '是否可自动化验证'],
            ['manual_steps', 'TEXT', '', 'JSON，人工指南步骤（不可验证时）'],
            ['execution_mode', 'TEXT', '', 'url_only / url_with_params / manual_guide'],
            ['verification_method', 'TEXT', '', 'direct / oob / manual'],
            ['input_schema', 'TEXT', '', 'JSON，url_with_params 的参数定义'],
        ],
        col_widths=[3.5, 2, 3, 7]
    )
    add_para(doc, '索引：idx_vuln_type(vuln_type)、idx_poc_type(poc_type)、idx_create_time(create_time DESC)、idx_verifiable(verifiable)。')

    add_heading(doc, '5.2.2 batch_tasks — 批量任务主表', level=3)
    add_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK AUTOINCREMENT', '主键'],
            ['task_type', 'TEXT', "DEFAULT 'poc_batch'", 'poc_batch / nuclei_batch'],
            ['mode', 'TEXT', 'NOT NULL', '执行模式（single_url_multi_poc 等）'],
            ['status', 'TEXT', "DEFAULT 'pending'", 'pending/running/completed/cancelled/failed'],
            ['total_items', 'INTEGER', 'DEFAULT 0', '笛卡尔积子任务总数'],
            ['completed_items', 'INTEGER', 'DEFAULT 0', '已完成子任务数'],
            ['success_items', 'INTEGER', 'DEFAULT 0', '成功数'],
            ['failed_items', 'INTEGER', 'DEFAULT 0', '失败数'],
            ['vulnerable_items', 'INTEGER', 'DEFAULT 0', '漏洞命中数'],
            ['concurrency', 'INTEGER', 'DEFAULT 3', '并发数'],
            ['config_json', 'TEXT', '', '任务创建参数快照'],
            ['error', 'TEXT', '', '任务级错误信息'],
            ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
            ['started_at', 'TIMESTAMP', '', '开始时间'],
            ['finished_at', 'TIMESTAMP', '', '结束时间'],
        ],
        col_widths=[3.5, 2, 3, 7]
    )
    add_para(doc, '索引：idx_batch_tasks_status(status)。')

    add_heading(doc, '5.2.3 batch_task_items — 批量任务子任务表', level=3)
    add_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK AUTOINCREMENT', '主键'],
            ['task_id', 'INTEGER', 'NOT NULL, FK→batch_tasks.id', '所属任务 ID'],
            ['poc_id', 'INTEGER', 'NOT NULL', 'POC ID；Nuclei 任务可为 0'],
            ['target_url', 'TEXT', 'NOT NULL', '目标 URL'],
            ['engine_type', 'TEXT', "DEFAULT 'poc'", 'poc / nuclei'],
            ['template_path', 'TEXT', '', 'Nuclei 模板相对路径'],
            ['status', 'TEXT', "DEFAULT 'pending'", 'pending/running/success/failed/cancelled'],
            ['result_json', 'TEXT', '', '结果摘要 JSON'],
            ['error', 'TEXT', '', '失败信息'],
            ['failure_category', 'TEXT', '', '失败大类（network/env/code/oob）'],
            ['failure_code', 'TEXT', '', '失败细分编码'],
            ['failure_stage', 'TEXT', '', '失败阶段（import/exec/verify）'],
            ['retryable', 'INTEGER', 'DEFAULT 0', '是否可重试'],
            ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
            ['started_at', 'TIMESTAMP', '', '开始时间'],
            ['finished_at', 'TIMESTAMP', '', '结束时间'],
        ],
        col_widths=[3.5, 2, 3.5, 6.5]
    )
    add_para(doc, '索引：idx_batch_task_items_task_id(task_id)、idx_batch_task_items_status(status)。')

    add_heading(doc, '5.3 文件存储结构', level=2)
    add_table(doc,
        ['路径', '内容', '命名规则 / 说明'],
        [
            ['pocs/poc_library.db', 'SQLite 数据库', '启动时自动创建'],
            ['pocs/python/', 'Python 版 POC 脚本', '<漏洞名>_<时间戳>_<hash>.py'],
            ['pocs/metadata/', '生成结果的原始 JSON 快照', '与 POC 脚本一一对应，便于回溯 LLM 原始输出'],
            ['pocs/nuclei/', 'Nuclei 模板库', '遵循 Nuclei 官方目录结构（http/cves、http/misconfiguration 等）'],
            ['pocs/llm_config.json', 'LLM 配置持久化文件', 'API Key 以脱敏方式展示，实际值仅在本地保存'],
            ['pocs/README.md', 'POC 库说明', ''],
        ],
        col_widths=[4, 5, 6.5]
    )

    doc.add_page_break()

    # ========== 6 接口设计 ==========
    add_heading(doc, '6 接口设计', level=1)

    add_heading(doc, '6.1 接口设计规范', level=2)
    add_bullet(doc, '统一前缀：所有业务接口挂载在 /api 下，静态资源通过 /、/css、/js、/assets 提供')
    add_bullet(doc, '报文格式：请求体与响应体默认使用 JSON，POC 生成与 Nuclei 流式扫描使用 SSE（text/event-stream）')
    add_bullet(doc, '状态码：200 正常；400 参数校验失败；404 资源不存在；500 内部异常，并在响应体 detail 中携带原因')
    add_bullet(doc, '通用字段：成功响应包含 success=true，以及具体业务字段；错误响应由 FastAPI 包装为 {"detail": "..."}')
    add_bullet(doc, '安全提示：涉及漏洞利用的响应统一附带 warning 字段，提醒用户仅用于授权测试')

    add_heading(doc, '6.2 主要接口列表', level=2)
    add_table(doc,
        ['方法', '路径', '功能说明'],
        [
            ['POST', '/api/generate-poc', '提交漏洞信息，返回 SSE 流式生成进度与最终 POC 结果'],
            ['GET', '/api/pocs/search', '按漏洞类型、关键字、是否可验证等条件搜索 POC'],
            ['GET', '/api/pocs/statistics', '获取 POC 库统计信息'],
            ['GET', '/api/pocs/vuln-types', '获取所有漏洞类型及其数量'],
            ['GET', '/api/pocs/{poc_id}', '获取 POC 详情'],
            ['GET', '/api/pocs/{poc_id}/code', '获取 POC 源代码内容'],
            ['POST', '/api/pocs/{poc_id}/execute', '对单个 POC 发起单目标验证'],
            ['DELETE', '/api/pocs/{poc_id}', '删除 POC 及其文件'],
            ['POST', '/api/batch-tasks', '创建批量检测任务'],
            ['GET', '/api/batch-tasks', '分页查询批量任务列表'],
            ['GET', '/api/batch-tasks/{task_id}', '获取批量任务详情'],
            ['GET', '/api/batch-tasks/{task_id}/items', '获取子任务列表（支持筛选）'],
            ['GET', '/api/batch-tasks/{task_id}/items/{item_id}/detail', '获取子任务详细结果'],
            ['POST', '/api/batch-tasks/{task_id}/cancel', '取消运行中的批量任务'],
            ['GET', '/api/batch-tasks/{task_id}/export', '导出任务报告（html/json/txt）'],
            ['GET', '/api/nuclei/status', '检查 Nuclei 可用性与模板总数'],
            ['GET', '/api/nuclei/folders', '获取模板目录结构'],
            ['GET', '/api/nuclei/templates', '分页获取模板列表'],
            ['GET', '/api/nuclei/template/content', '查看模板 YAML 内容'],
            ['POST', '/api/nuclei/scan', '执行即时 Nuclei 扫描'],
            ['POST', '/api/nuclei/scan/stream', '以 SSE 形式执行 Nuclei 扫描'],
            ['POST', '/api/nuclei/tasks', '创建 Nuclei 任务化扫描'],
            ['POST', '/api/nuclei/cache/clear', '清除模板解析缓存'],
            ['GET/POST', '/api/config/llm', '查询/更新主生成模型配置'],
            ['GET/POST', '/api/config/review-llm', '查询/更新二次审核模型配置'],
            ['GET/POST', '/api/config/oob', '查询/更新 OOB 配置'],
            ['GET/POST', '/api/config/asset-sources', '查询/更新空间测绘平台配置'],
            ['POST', '/api/asset-sources/import', '从空间测绘平台导入目标 URL'],
        ],
        col_widths=[1.8, 6.5, 7.2]
    )

    add_heading(doc, '6.3 典型报文示例', level=2)
    add_para(doc, '（1）POC 生成请求：')
    add_code(doc, '''POST /api/generate-poc
Content-Type: application/json

{
  "vulnerability_info": "目标网站存在 SQL 注入漏洞，位于登录页的 username 参数",
  "enable_second_review": true
}''')
    add_para(doc, '响应为 SSE 事件流，关键事件包含：')
    add_code(doc, '''data: {"type":"status","step":1,"message":"正在使用 gpt-4 生成 POC..."}
data: {"type":"status","step":2,"message":"正在二次审核..."}
data: {"type":"result","data":{"success":true,"vulnerability_type":"SQL注入",
  "poc_code":"import requests\\n...","execution_mode":"url_only",
  "verification_method":"direct","dependency_check":{"ok":true}}}
data: [DONE]''')

    add_para(doc, '（2）批量任务创建：')
    add_code(doc, '''POST /api/batch-tasks
Content-Type: application/json

{
  "target_urls": ["http://a.com", "http://b.com"],
  "poc_ids": [101, 102, 103],
  "concurrency": 3
}''')

    doc.add_page_break()

    # ========== 7 关键流程设计 ==========
    add_heading(doc, '7 关键流程设计', level=1)

    add_heading(doc, '7.1 POC 生成流程', level=2)
    add_para(doc, 'POC 生成是系统最核心的流程，分为"初稿生成 → 二次审核（可选）→ 依赖预检 → 入库保存"四个阶段。每个阶段完成后通过 SSE 向前端推送状态，用户可实时看到生成进度。')
    add_code(doc, '''
[前端] 提交漏洞信息
   │
   ▼
[routes.generate_poc] 打开 SSE
   │
   ▼
[llm_service._build_prompt] 组装系统/用户 Prompt
   │
   ▼
[llm_service.generate_initial_poc]  → LLM API → 解析 JSON
   │
   ├─── 启用二次审核 ──▶ [review_generated_poc] → LLM API
   │                                        ▼
   │                              合并审核结果
   ▼
[poc_library_service.check_poc_dependencies] 依赖静态检查
   │
   ▼
[poc_library_service.save_poc] 写 SQLite + 落盘 .py 文件
   │
   ▼
推送 result 事件 → [DONE]
'''.strip('\n'))

    add_heading(doc, '7.2 POC 单条验证流程', level=2)
    add_para(doc, '用户在 POC 库选中某个 POC 并输入目标 URL 后，系统根据 execution_mode 分发执行：')
    add_bullet(doc, 'url_only：仅传入 URL，直接调用 POC 的 scan(url) 函数')
    add_bullet(doc, 'url_with_params：根据 input_schema 让用户填写附加参数（如 Cookie、userID），再调用 scan(url, **kwargs)')
    add_bullet(doc, 'manual_guide：不可自动化验证，前端展示 manual_steps 的逐步操作指引')
    add_para(doc, '在可验证分支中，系统会临时给 requests 会话打上超时补丁、动态 import POC 脚本模块、捕获异常并交由 failure_classifier 分类，最终返回 success、vulnerable、classification 等结构化字段。')

    add_heading(doc, '7.3 批量检测流程', level=2)
    add_code(doc, '''
[create_task]
   │ 去重校验 + 写 batch_tasks / batch_task_items
   ▼
[start_task] 启动后台线程
   │ 取出 pending 子任务 → 线程池并发
   ▼
[_execute_task_item]
   │   POC 任务：poc_library_service.execute_poc()
   │   Nuclei 任务：nuclei_service.scan_single()
   ▼
[_store_item_result] 记录单子任务结果
   │
   ▼
[_refresh_task_stats] 实时刷新任务总进度
   │
   ▼
[_finalize_task] 全部完成或被取消后落库
'''.strip('\n'))
    add_para(doc, '整个过程支持用户通过前端轮询 /api/batch-tasks/{id} 获取任务状态，并可随时通过取消接口置位 cancel_event 停止尚未开始的子任务。')

    add_heading(doc, '7.4 OOB 外带验证流程', level=2)
    add_para(doc, '针对无回显类漏洞，POC 可以通过 oob_service 获取一次性 flag 和回连地址，将其注入 payload 发给目标；目标如果触发漏洞，会主动向回连地址发起 DNS 或 HTTP 请求；POC 在执行后轮询 OOB 平台确认是否命中，从而判断漏洞是否存在。整个过程对上层 API 透明，只需在 POC 内显式声明 verification_method = "oob"。')

    add_heading(doc, '7.5 Nuclei 扫描流程', level=2)
    add_para(doc, 'Nuclei 扫描区分即时扫描与任务化扫描两条路径：即时扫描由 nuclei_service 调起子进程并实时解析 stdout，以 SSE 形式回传 finding；任务化扫描则走与 POC 批量任务相同的调度框架，进入统一的 batch_tasks 体系，便于在检测记录页一并查看结果。')

    doc.add_page_break()

    # ========== 8 安全设计 ==========
    add_heading(doc, '8 安全设计', level=1)

    add_heading(doc, '8.1 使用授权与合规', level=2)
    add_para(doc, '系统在启动时控制台即打印安全警告，所有 POC 生成响应统一附带 warning 字段提示"仅用于授权的安全测试和研究目的"。系统默认仅监听 127.0.0.1，阻止局域网其他设备访问，避免被他人滥用。如需对外提供服务，需由使用者自行修改 config.py 并承担相应法律责任。')

    add_heading(doc, '8.2 敏感信息保护', level=2)
    add_bullet(doc, 'LLM API Key、OOB Token、空间测绘平台 Token 等敏感配置保存在本地 JSON 文件，不会随代码仓库提交；前端接口返回时统一做脱敏（仅保留前后 4 位）')
    add_bullet(doc, 'Interactsh 通信采用 RSA + AES 混合加密，防止监听方推测回连内容')
    add_bullet(doc, '日志中不打印完整 API Key 和原始 payload 中的账号密码字段')

    add_heading(doc, '8.3 输入与执行安全', level=2)
    add_bullet(doc, 'Pydantic 模型对所有入参做类型与长度校验，拒绝非预期字段')
    add_bullet(doc, '批量任务对 URL 数量与 POC 数量均设上限（MAX_URLS、MAX_POCS），防止资源耗尽')
    add_bullet(doc, 'POC 执行时 requests 会话强制超时，避免恶意目标通过慢响应拖垮后端')
    add_bullet(doc, 'LLM 返回结果经 dependency_checker 静态扫描，过滤潜在风险 import；执行异常统一捕获并归因，避免堆栈直接暴露给前端')
    add_bullet(doc, 'CORS 当前为开发期宽松策略，正式部署前需替换为固定域名白名单')

    add_heading(doc, '8.4 日志与审计', level=2)
    add_para(doc, '使用 RotatingFileHandler 输出到 api_server.log，按 10MB 切分、保留 5 份历史；同时通过自定义中间件对每次 HTTP 请求打印方法、URL、客户端 IP、耗时与状态码，便于追溯异常行为与问题定位。')

    # ========== 9 部署 ==========
    add_heading(doc, '9 运行部署设计', level=1)

    add_heading(doc, '9.1 部署模式', level=2)
    add_para(doc, '系统提供两种部署模式：开发模式和便携模式。开发模式适合持续开发与调试，依赖开发者自行安装 Python 3.11 与 requirements.txt；便携模式将源码、虚拟环境、Nuclei 可执行文件与启动脚本一并打包到 release/ 目录，双击 start.bat 即可在未安装 Python 的机器上运行，适合课程答辩与非技术用户演示。')

    add_heading(doc, '9.2 启动步骤', level=2)
    add_code(doc, '''# 开发模式
pip install -r requirements.txt
python main.py
# 浏览器访问 http://127.0.0.1:8000/

# 便携模式
release/start.bat
# 自动调起内置 Python 与 Uvicorn''')

    add_heading(doc, '9.3 外部依赖配置', level=2)
    add_table(doc,
        ['外部依赖', '配置方式', '是否必须'],
        [
            ['LLM 服务', '前端"系统配置 → LLM"填写 API Key / 模型 ID / Base URL', '必需（否则无法生成 POC）'],
            ['二次审核模型', '同上，入口为"二次审核 LLM"', '可选'],
            ['OOB 服务', '前端"系统配置 → OOB"选择 Interactsh 或 CEye 并填写凭证', '可选，仅无回显类漏洞需要'],
            ['Nuclei', '使用项目根目录自带 nuclei.exe 或在环境变量 PATH 中配置', '可选'],
            ['空间测绘平台', '前端"系统配置 → 空间测绘"填写 FOFA / Hunter / Quake 凭证', '可选，仅批量导入目标时需要'],
        ],
        col_widths=[3, 8, 4.5]
    )

    doc.add_page_break()

    # ========== 10 技术难点 ==========
    add_heading(doc, '10 技术难点与解决方案', level=1)

    add_table(doc,
        ['技术难点', '难点分析', '解决方案'],
        [
            ['LLM 输出不稳定',
             '大模型返回可能包含非标准 JSON、乱码、控制字符或冗余内容，直接解析会失败；生成的 POC 也可能缺少依赖或使用不存在的 API',
             '① 双阶段 Prompt（初稿 + 二次审核）；② _parse_llm_json_response 对控制字符与转义错误做鲁棒修复；③ dependency_checker 静态扫描 import；④ 失败时保留初稿而不是直接报错'],
            ['多执行模式兼容',
             '直接验证、带参验证、人工指南、OOB 验证四种形态差异较大，简单用 if/else 难以维护',
             '引入 execution_mode 与 verification_method 两个枚举字段，前后端按模式分支渲染；POC 脚本统一暴露 scan(url, **kwargs) 入口'],
            ['批量任务调度',
             'URL × POC 的笛卡尔积任务数量可能达到数千个，同步执行会阻塞 Web 服务；同时需要支持取消和进度查询',
             '两级任务模型 batch_tasks + batch_task_items；使用线程池执行子任务；cancel_event 协同取消；实时刷新统计到主表以便前端轮询'],
            ['失败结果可解释性',
             '仅返回"失败"不利于调试；不同失败类型（网络 / 环境 / 代码 / OOB）的处置方式差异很大',
             'failure_classifier 结合异常类型、堆栈信息、HTTP 响应等判定 failure_category / failure_code / failure_stage，前端按类别呈现'],
            ['Nuclei 模板海量解析',
             '官方模板库规模上万，完整加载 YAML 耗时较长',
             '_parse_template_fast 只解析头部字段；模板目录与列表结果缓存，提供 cache/clear 接口按需重建'],
            ['跨平台演示',
             '非技术用户环境不稳定，现场安装 Python 成本高',
             '便携版打包：内置解释器与依赖；start.bat 一键启动；日志轮转避免磁盘占用过大'],
            ['前端无构建工具',
             '为保持演示简单，不引入 webpack/vite；但仍需管理多页面、多模块的状态',
             '按功能拆分 app.js / library.js / nuclei.js，使用 Fetch + EventSource 原生 API；开发期静态资源禁用缓存'],
        ],
        col_widths=[3, 5.5, 7]
    )

    # ========== 11 附录 ==========
    add_heading(doc, '11 附录', level=1)

    add_heading(doc, '11.1 执行模式取值说明', level=2)
    add_table(doc,
        ['execution_mode', 'verification_method', '适用场景'],
        [
            ['url_only', 'direct', '仅需目标 URL 即可验证（如反射型 XSS、开源组件版本指纹）'],
            ['url_only', 'oob', '无回显漏洞，依赖外带通道（如 SSRF、盲注、命令执行）'],
            ['url_with_params', 'direct', '需要额外参数（Cookie、用户名、参数名等）才能验证'],
            ['url_with_params', 'oob', '带参 + 外带，常见于复杂注入链'],
            ['manual_guide', 'manual', '不可自动化，只能依照步骤手动验证（业务逻辑漏洞、多步攻击链）'],
        ],
        col_widths=[3.5, 3.5, 8.5]
    )

    add_heading(doc, '11.2 Python 依赖清单（requirements.txt）', level=2)
    add_table(doc,
        ['依赖', '版本约束', '用途'],
        [
            ['fastapi', '>=0.110.0', 'Web 框架'],
            ['uvicorn[standard]', '>=0.27.0', 'ASGI 服务器'],
            ['pydantic', '>=2.6.0', '数据校验'],
            ['openai', '>=1.0.0', '兼容 OpenAI 协议的 LLM 客户端'],
            ['PyYAML', '>=6.0.1', 'Nuclei 模板与 Prompt 配置解析'],
            ['dnspython', '>=2.4.0', 'OOB DNS 查询'],
            ['requests', '>=2.31.0', 'POC 内部 HTTP 调用'],
            ['pycryptodome', '>=3.20.0', 'Interactsh AES 解密'],
            ['itsdangerous', '>=2.2.0', '敏感配置签名/加密'],
            ['Flask', '>=3.0.0', '兼容少量旧版工具接口'],
        ],
        col_widths=[3.5, 3, 8.5]
    )

    add_heading(doc, '11.3 文档维护', level=2)
    add_para(doc, '本文档随项目功能演进同步更新，后续主要维护点包括：① 新增 API 或 Service 时同步更新第 4、6 节；② 数据库字段变更同步更新第 5 节；③ 重要技术方案变更同步更新第 10 节；④ 界面改版时同步补充前端模块章节。')

    # 保存
    doc.save(str(OUTPUT_PATH))
    print(f'文档已生成: {OUTPUT_PATH}')
    print(f'文件大小: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB')


if __name__ == '__main__':
    build()
